"""Build docs/DAS_Sentinel_User_Guide.docx.

Content is a flat list of (kind, payload) blocks; a small renderer maps each kind
onto python-docx. The API reference table is generated from the live route map
(caps.json) so it cannot drift from the code.
"""

import json
import os
import pathlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRATCH = pathlib.Path(os.environ["SCRATCH"])
REPO = pathlib.Path(os.environ["REPO"])
OUT = REPO / "docs" / "DAS_Sentinel_User_Guide.docx"

CAPS = json.loads((SCRATCH / "caps.json").read_text())

ACCENT = RGBColor(0x1F, 0x3B, 0x57)
MUTED = RGBColor(0x55, 0x5F, 0x67)


# ─────────────────────────────── renderer ────────────────────────────────


def styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before in (
        ("Title", 30, ACCENT, 0),
        ("Heading 1", 18, ACCENT, 20),
        ("Heading 2", 13.5, ACCENT, 14),
        ("Heading 3", 11.5, ACCENT, 10),
        ("Heading 4", 10.5, MUTED, 8),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = name != "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def add_toc(doc):
    """A real Word TOC field. `updateFields` below makes Word populate it on open."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    inner = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose “Update field” if the index is blank."
    inner.append(t)
    fld.append(inner)
    p._p.append(fld)


def update_fields_on_open(doc):
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def runs(paragraph, text):
    """Inline markup: **bold**, *italic*, `mono`. ** is matched first."""
    import re

    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def render(doc, blocks):
    for kind, payload in blocks:
        if kind == "h1":
            doc.add_heading(payload, 1)
        elif kind == "h2":
            doc.add_heading(payload, 2)
        elif kind == "h3":
            doc.add_heading(payload, 3)
        elif kind == "h4":
            doc.add_heading(payload, 4)
        elif kind == "p":
            runs(doc.add_paragraph(), payload)
        elif kind == "b":
            for line in payload:
                runs(doc.add_paragraph(style="List Bullet"), line)
        elif kind == "n":
            for line in payload:
                runs(doc.add_paragraph(style="List Number"), line)
        elif kind == "note":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run("⚠  ")
            r.bold = True
            runs(p, payload)
            for r in p.runs:
                r.font.size = Pt(10)
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(payload)
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
        elif kind == "table":
            headers, rows, widths = payload
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
            for i, h in enumerate(headers):
                cell = t.rows[0].cells[i]
                cell.text = ""
                runs(cell.paragraphs[0], f"**{h}**")
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    runs(cells[i].paragraphs[0], str(val))
            for row in t.rows:
                for i, cell in enumerate(row.cells):
                    cell.width = Inches(widths[i])
                    for p in cell.paragraphs:
                        p.paragraph_format.space_after = Pt(2)
                        for r in p.runs:
                            r.font.size = Pt(8.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(kind)


# ──────────────────────────────── content ────────────────────────────────

W3 = (2.0, 2.2, 2.4)
W2 = (2.4, 4.2)

ROLE_MATRIX = [
    ["Manage users (create, role, deactivate, reset)", "Yes", "No", "No", "No"],
    ["Register / remove AI models", "Yes", "No", "No", "No"],
    ["Create, edit, delete engagements", "Yes", "Yes", "No", "No"],
    ["Change engagement status", "Yes", "Yes", "No", "No"],
    ["Edit scope (allow / deny)", "Yes", "Yes", "No", "No"],
    ["Accept the Rules of Engagement", "Yes", "Yes", "No", "No"],
    ["Add / edit / delete targets", "Yes", "Yes", "No", "No"],
    ["Upload a source archive", "Yes", "Yes", "No", "No"],
    ["Manage credentials", "Yes", "Yes", "No", "No"],
    ["Launch and cancel scans", "Yes", "Yes", "No", "No"],
    ["Request a high-risk approval gate", "Yes", "Yes", "No", "No"],
    ["Approve / deny / revoke a gate", "Yes", "No", "Yes", "No"],
    ["Score CVSS, map compliance, generate remediation", "Yes", "Yes", "Yes", "No"],
    ["Generate, edit, finalize, export reports", "Yes", "Yes", "Yes", "No"],
    ["View the audit log", "Yes", "No", "Yes", "No"],
    ["View engagements, targets, scans, findings, reports", "Yes", "Yes", "Yes", "Yes"],
    ["Own account settings (/profile)", "Yes", "Yes", "Yes", "Yes"],
]

BLOCK_REASONS = [
    [
        "engagement_inactive",
        "409",
        "The engagement is not Active.",
        "Open the engagement and click **Activate** in the Status card.",
    ],
    [
        "roe_not_accepted",
        "403",
        "No Rules of Engagement acknowledgement exists.",
        "Tick the acknowledgement box and click **Accept Rules of Engagement**.",
    ],
    [
        "roe_stale",
        "403",
        "Scope changed after the ROE was accepted, so the content hash no longer matches.",
        "Re-accept the ROE.",
    ],
    [
        "roe_terms_mismatch",
        "409",
        "An engagement term (window, rate limit, max intensity) changed after acceptance.",
        "Re-accept the ROE.",
    ],
    [
        "outside_test_window",
        "403",
        "Now is outside the authorized window — **including when no window is set at all**.",
        "Edit the engagement so the window covers now, then re-accept the ROE.",
    ],
    [
        "scope_violation",
        "403",
        "The target matches no allow rule, or matches a deny rule. Deny always wins.",
        "Add an allow rule that covers the target, or remove the deny rule.",
    ],
    [
        "intensity_not_authorized",
        "422",
        "The server-derived effective intensity exceeds the engagement ceiling.",
        "Lower the intensity or raise **Maximum intensity** on the engagement (then re-accept the ROE).",
    ],
    [
        "high_risk_not_approved",
        "403",
        "The operation is high-risk and has no valid approval gate.",
        "Request a gate and have a Reviewer or Admin approve it (API only — see §18).",
    ],
    [
        "ssrf_ip_blocked",
        "403",
        "The target resolves to a loopback, private, or cloud-metadata address.",
        "Point the target at a routable, authorized host. Scope allow-listing never overrides this.",
    ],
]

FRAMEWORKS = [
    ["OWASP Top 10 for LLM Applications", "2025", "10", "LLM01 – LLM10"],
    ["OWASP Top 10 for Agentic Applications", "2026", "10", "ASI01 – ASI10"],
    ["OWASP Web Security Testing Guide", "4.2", "12", "WSTG-INFO … WSTG-APIT"],
    ["NIST AI Risk Management Framework", "1.0 (AI 100-1)", "19", "GOVERN / MAP / MEASURE / MANAGE"],
    ["NIST AI RMF Generative AI Profile", "AI 600-1 (July 2024)", "12", "cbrn … value-chain"],
    ["NIST SP 800-53", "Rev 5.2.0", "16", "AC-3, AC-6, AU-2 … SI-11"],
    ["NIST SP 800-115", "September 2008", "6", "SP800-115-PLANNING … -POST"],
]

BLOCKS = []
A = BLOCKS.append


def sec(title):
    A(("h1", title))


# ── 1
A(("pagebreak", None))
sec("1  About this guide")
A(
    (
        "p",
        "This guide explains how to operate **DAS Sentinel** end to end: every screen, "
        "every control on every screen, every field, every error you can hit, and the "
        "features that currently exist only on the API. It is written for the people who "
        "actually use the platform — security testers, engineers, and compliance reviewers.",
    )
)
A(("h2", "1.1  How to read it"))
A(
    (
        "b",
        [
            "Sections 5 to 8 are setup and orientation — read them once.",
            "Sections 9 to 16 follow the working order of a real assessment: engagement → scope → ROE → targets → scan.",
            "Sections 17 to 25 cover results, safety, and oversight: blocked launches, findings, scoring, compliance, reports, audit, health.",
            "Section 26 lists the features that have no screen yet and how to drive them over the API.",
            "Section 27 is the complete endpoint reference. Section 28 lists the current limitations honestly. Section 29 is troubleshooting.",
        ],
    )
)
A(("h2", "1.2  Conventions"))
A(
    (
        "b",
        [
            "`Monospace` marks something you type or something the system returns verbatim — a URL, a field value, an API path, a machine-readable reason code.",
            "**Bold** marks a control you click: a button, a link, a menu item, a tab.",
            "A ⚠ paragraph is a safety or data-loss warning. Read those before acting.",
            "“Sidebar → Item” means click that item in the left navigation.",
        ],
    )
)
A(("h2", "1.3  Accuracy and verification"))
A(
    (
        "p",
        "Every behaviour in this guide was checked against the running application and its "
        "source, not against design documents. The verification record — what was executed, "
        "what passed, and which external standards were confirmed against their publishers — "
        "is Appendix B (§31). Where the product is incomplete, this guide says so plainly in "
        "§28 rather than describing an intention as a feature.",
    )
)

# ── 2
sec("2  What DAS Sentinel does")
A(
    (
        "p",
        "DAS Sentinel is an AI security-testing and automated penetration-testing platform "
        "for **authorized defensive security assessments** of web applications, APIs, source "
        "code, and AI/LLM applications. It turns an approved engagement scope into "
        "evidence-backed, compliance-mapped, prioritized, report-ready findings.",
    )
)
A(
    (
        "p",
        "The value is not the scanning. The value is that every finding carries its evidence, "
        "its provenance, a CVSS score, and a mapping to a published control — so the output is "
        "defensible to an auditor.",
    )
)
A(("h2", "2.1  What it will not do"))
A(
    (
        "p",
        "These are hard product rules, enforced in the service layer rather than hidden in the UI. "
        "You cannot switch them off:",
    )
)
A(
    (
        "n",
        [
            "**No scan without authorization.** A saved engagement, a defined scope, and an accepted ROE are all required before any test runs.",
            "**Every target is re-checked against the allowlist and blocklist** immediately before execution — once at launch, again in the worker. Deny rules always win.",
            "**Safe and non-destructive by default.** Higher intensity must be explicitly selected *and* authorized by the engagement ceiling.",
            "**High-risk actions need an explicit approval gate** — exploit validation, brute force, large-scale crawling, data-modifying payloads. There is no auto-exploitation.",
            "**No offensive capability.** No stealth, evasion, persistence, credential harvesting, exfiltration, or denial-of-service.",
            "**The AI is never the source of truth.** Every finding cites concrete evidence, and an AI-generated finding is labelled as a draft until a human validates it.",
            "**Redaction before egress.** Nothing reaches a hosted model unless the engagement allows hosted models, and redaction runs first. If redaction errors, egress is blocked.",
            "**Everything is audited**, append-only — including every blocked attempt.",
            "**Emergency stop always works.** A running scan is cancellable, and cancellation kills the underlying process tree.",
        ],
    )
)
A(
    (
        "note",
        "Never point a scanner or an AI/LLM suite at a system you are not authorized to test. "
        "Use the `sandbox/` mock apps, an intentionally-vulnerable lab, or a system you own. "
        "Everything you do is attributed to your account and audited.",
    )
)

# ── 3
sec("3  Concepts and vocabulary")
A(
    (
        "table",
        (
            ["Term", "What it means in this product"],
            [
                [
                    "Engagement",
                    "The authorization container. Holds the client/system name, the test window, "
                    "the rate limit, the intensity ceiling, the hosted-model flag, and the contacts. "
                    "Nothing exists outside an engagement.",
                ],
                [
                    "Scope item",
                    "One allow or deny rule. Five matcher types: URL, domain, IP/CIDR, API base, "
                    "repository. Deny always beats allow.",
                ],
                [
                    "ROE (Rules of Engagement)",
                    "The acknowledgement that authorizes testing. Accepting it writes an immutable "
                    "row with a SHA-256 content hash over the ROE text, the scope snapshot, and the "
                    "engagement terms. Change any of those and it must be re-accepted.",
                ],
                [
                    "Target",
                    "One system under test, typed (web app, REST API, source repo, AI chatbot, …). "
                    "Its type fixes which scanners and suites can run against it.",
                ],
                [
                    "Scan",
                    "One execution record: either a set of AI/LLM suites or a set of code/web "
                    "scanners against one target, never both in one scan.",
                ],
                [
                    "Intensity",
                    "How aggressive a run is: `passive`, `safe_active`, `authenticated_active`, "
                    "`high_risk`. The server derives the effective intensity — a caller cannot "
                    "declare a higher one than its operation actually warrants.",
                ],
                [
                    "Finding",
                    "A normalized result with severity, evidence, provenance, status, an optional "
                    "OWASP reference, CVSS scores, and compliance mappings.",
                ],
                [
                    "Provenance",
                    "Where a finding came from and how far it has been reviewed: `automated`, "
                    "`ai_generated`, `validated`, `manually_overridden`. The first two are never "
                    "presented as verified.",
                ],
                [
                    "Evidence",
                    "The raw artefact behind a finding — scanner output, an HTTP transcript, an LLM "
                    "conversation, a source archive — stored content-hashed in the object store and "
                    "re-verified by SHA-256 every time it is read.",
                ],
                [
                    "Approval gate",
                    "A single-use authorization for one high-risk operation, bound to the current "
                    "ROE acknowledgement and to a digest of that exact operation.",
                ],
                [
                    "Capability",
                    "The unit of permission. Roles hold capabilities; routes require capabilities. "
                    "See §4.",
                ],
            ],
            (1.5, 5.1),
        ),
    )
)

# ── 4
sec("4  Access, roles and permissions")
A(("h2", "4.1  The four roles"))
A(
    (
        "b",
        [
            "**Admin** — everything: user management, AI model registration, credentials, engagements, scope, ROE, targets, scans, approving gates, validation, reports, audit log.",
            "**Tester** — the operator role. Manages engagements, scope, ROE, targets and credentials; launches and cancels scans; requests high-risk gates; validates findings; generates and exports reports. No user management, no AI model registration, no audit log, and **cannot approve a gate they requested**.",
            "**Reviewer** — the oversight role. Views everything, approves/denies/revokes high-risk gates, validates findings (CVSS + compliance mapping), exports reports, reads the audit log. Cannot create engagements, launch scans, or touch credentials.",
            "**Read only** — views engagements, scope, ROE, targets, scans, findings, reports, AI models and health. Changes nothing, cannot export reports, and has no access to credentials, users, or the audit log.",
        ],
    )
)
A(("h2", "4.2  Permission matrix"))
A(
    (
        "table",
        (
            ["Action", "Admin", "Tester", "Reviewer", "Read only"],
            ROLE_MATRIX,
            (3.4, 0.8, 0.8, 0.85, 0.95),
        ),
    )
)
A(("h2", "4.3  How enforcement actually works"))
A(
    (
        "p",
        "Sidebar items and buttons are hidden by role for convenience. **The API is the "
        "enforcement.** Reaching a forbidden action directly by URL returns `403` with a role "
        "message, and nothing changes. The capability table lives in one place in the code "
        "(`apps/api/app/core/deps.py`), so a route cannot drift from the matrix above.",
    )
)
A(
    (
        "p",
        "Access control across engagements returns `404`, never `403` and never data — so you "
        "cannot even confirm that another team's engagement exists.",
    )
)
A(("h2", "4.4  Session rules"))
A(
    (
        "b",
        [
            "Sessions are **opaque server-side tokens**, not JWTs, so revocation is instant — every request re-validates against the store.",
            "The cookie is `__Host-das_session`: HttpOnly, Secure, SameSite=Strict. A session identifier never appears in a URL.",
            "**15-minute sliding idle timeout** and an **8-hour absolute cap**. Pause for more than 15 minutes and your next click lands on `/login?expired=1`.",
            "A role change, a deactivation, an admin password reset, or **Sign out everywhere** revokes sessions immediately, in both the cache and the database.",
            "Every state-changing request needs the `X-CSRF-Token` header as well as the cookie. A cookie alone is refused with `403`.",
        ],
    )
)
A(
    (
        "note",
        "A **self-service password change does not revoke your other sessions**. If you think an "
        "account is compromised, change the password *and* use **Sign out everywhere**.",
    )
)

# ── 5
sec("5  Getting the platform running")
A(
    (
        "p",
        "Skip this section if someone has already given you a URL and an account. It is here for "
        "whoever stands the stack up.",
    )
)
A(("h2", "5.1  Start the stack"))
A(("code", "cp .env.example .env      # then fill in the values\ndocker compose up -d"))
A(
    (
        "p",
        "The `migrate` service applies the database schema automatically before `api` starts. "
        "Wait until `docker compose ps` shows `api`, `web`, `proxy`, `postgres` and `valkey` as "
        "healthy. The application is then served by the Caddy reverse proxy at "
        "`https://localhost` — in development the TLS certificate is self-signed, so accept the "
        "browser warning once.",
    )
)
A(("h2", "5.2  Seed the compliance knowledge base"))
A(
    (
        "p",
        "Compliance mapping needs its control catalogue loaded. Without it, the Framework "
        "dropdown on a finding is empty:",
    )
)
A(("code", "python apps/api/scripts/seed_compliance.py"))
A(("h2", "5.3  Create the first administrator"))
A(
    (
        "p",
        "The platform ships with **no users**. Create the first admin inside the compose network:",
    )
)
A(
    (
        "code",
        'docker compose run --rm --no-deps \\\n'
        '  -v "$PWD/apps/api/scripts:/app/scripts:ro" --entrypoint sh api \\\n'
        '  -c "cd /app && PYTHONPATH=/app uv run --no-sync python scripts/seed_e2e_user.py"',
    )
)
A(
    (
        "p",
        "Sign in as that admin, then create the real accounts from **Sidebar → Users** (§7 and "
        "§18 of this guide cover the flow from both sides).",
    )
)
A(("h2", "5.4  Browser requirements"))
A(
    (
        "b",
        [
            "Current Chrome, Edge, or Firefox. Use a clean profile or a private window so no stale session cookie interferes.",
            "The application makes **no third-party network calls** — fonts are self-hosted and `connect-src` is `'self'`. It runs air-gapped by design.",
            "Every page carries a Content-Security-Policy with a per-request nonce. If you see CSP violations in the console during normal use, that is a bug worth reporting.",
        ],
    )
)

# ── 6
sec("6  Signing in")
A(("h2", "6.1  The sign-in page"))
A(
    (
        "p",
        "Browse to `https://localhost`. If you have no session you are redirected to `/login`. "
        "The tab title reads *Sign in — DAS Sentinel*; the page shows a shield mark, the heading "
        "**DAS Sentinel**, the subtitle *AI-based Automated Security Testing*, **Email** and "
        "**Password** fields, a **Sign in** button, and the footer *DAS Sentinel · authorized "
        "defensive assessments only*.",
    )
)
A(("h3", "Steps"))
A(
    (
        "n",
        [
            "Enter your email address.",
            "Enter your password.",
            "Click **Sign in**. The button reads *Signing in…* while it works.",
            "On success the whole application reloads and you land on the **Dashboard**. A round initials avatar appears at the top-right of the header.",
        ],
    )
)
A(("h2", "6.2  When sign-in fails"))
A(
    (
        "table",
        (
            ["What you did", "What you see", "Why"],
            [
                ["Wrong password", "*Invalid email or password.*", "Credential rejected."],
                [
                    "Email with no account",
                    "*Invalid email or password.* — identical",
                    "Deliberate: the form must not reveal whether an address exists.",
                ],
                [
                    "Submitted both fields blank",
                    "*Invalid email or password.*",
                    "A malformed credential reads the same as a wrong one, so it is not an oracle.",
                ],
                [
                    "Correct credentials on a deactivated account",
                    "*Invalid email or password.*",
                    "Deactivated accounts are indistinguishable from wrong credentials.",
                ],
                [
                    "Too many failures",
                    "*Sign-in failed — the API is unreachable. Try again.*",
                    "You are being rate-limited (`429`). **5 failures per email** and **30 per IP** "
                    "in a rolling 15-minute window. Waiting clears it.",
                ],
                [
                    "MFA is enabled on the account",
                    "*Invalid email or password.*",
                    "A known gap. The API answers `mfa_required`, but the form has no code field, "
                    "so an MFA-enabled account cannot finish sign-in in the browser. See §28.",
                ],
            ],
            (1.9, 2.3, 2.4),
        ),
    )
)
A(("h2", "6.3  The expired-session banner"))
A(
    (
        "p",
        "If your session idled out or was revoked, you land on `/login?expired=1` with an amber "
        "banner: *Your session has expired or was revoked. Sign in again.* A visitor who never "
        "had a session gets the plain `/login` with no banner.",
    )
)
A(("h2", "6.4  First sign-in with a temporary password"))
A(
    (
        "p",
        "An administrator never chooses your password. They create the account and the platform "
        "generates a **one-time temporary password**, shown to them once. Your first sign-in with "
        "it succeeds but lands on `/set-password`, not the Dashboard.",
    )
)
A(
    (
        "p",
        "That page has **no sidebar and no account menu** — there is nothing to click past it. It "
        "shows the heading **Set your password**, the note *You signed in with a temporary "
        "password. Choose a permanent one to continue.*, **New password** and **Confirm "
        "password** fields, and a **Set password** button.",
    )
)
A(("h3", "Steps"))
A(
    (
        "n",
        [
            "Type your chosen permanent password in **New password**.",
            "Type the same value in **Confirm password**.",
            "Click **Set password**. The button reads *Saving…*, then you land on the Dashboard with the same session still valid — no second sign-in needed.",
        ],
    )
)
A(("h3", "Password rules"))
A(
    (
        "table",
        (
            ["If you enter…", "You get"],
            [
                [
                    "Fewer than 12 characters",
                    "*Password must be at least 12 characters.* — blocked in the browser, no request sent.",
                ],
                ["A confirmation that differs", "*Passwords do not match.* — nothing is saved."],
                [
                    "A known-breached or common password (e.g. `Password1234!`)",
                    "*password appears in a known-breach/common-password list; choose another* — refused by the server (`422`).",
                ],
                [
                    "A valid password",
                    "Saved. The temporary password stops working immediately, and you are never sent to `/set-password` again.",
                ],
            ],
            (2.5, 4.1),
        ),
    )
)
A(
    (
        "note",
        "The forced change is **enforced, not advisory**. While your account is in forced-change "
        "state, every page bounces back to `/set-password` and every API endpoint except "
        "`GET /auth/me`, `POST /auth/me/password` and the two logout routes returns "
        "`403 password_change_required`. A shared temporary password cannot be used as a working "
        "account.",
    )
)
A(("h2", "6.5  Signing out"))
A(
    (
        "b",
        [
            "**Sign out** — click the initials avatar, then **Sign out**. Your session is revoked and you land on `/login`. Pressing Back does not restore an authenticated page.",
            "**Sign out everywhere** — same menu. A dialog asks *Sign out of every session on every device?* Confirming revokes **all** your sessions; the others are bounced to `/login` on their next request. Cancelling does nothing.",
        ],
    )
)

# ── 7
sec("7  Your account")
A(
    (
        "p",
        "Every signed-in role has account settings, **including Read only**. They live behind the "
        "initials avatar at the top-right of the header — not in the sidebar.",
    )
)
A(("h2", "7.1  The account menu"))
A(
    (
        "p",
        "Click the avatar. The panel shows your display name in bold, your email muted beneath "
        "it, a badge with your role (**Admin** / **Tester** / **Reviewer** / **Read only**), then "
        "**Account settings**, **Sign out**, and **Sign out everywhere**.",
    )
)
A(("h2", "7.2  Account settings"))
A(
    (
        "p",
        "**Account menu → Account settings** opens `/profile`, titled *Account settings* with the "
        "note *Update your profile and password. Changes are attributed and audited.* Two cards: "
        "**Profile** (Name, Email, Phone number) and **Password** (Current, New, Confirm).",
    )
)
A(("h3", "7.2.1  Change your display name"))
A(
    (
        "n",
        [
            "Edit **Name**.",
            "Click **Save profile**. A green *Profile saved.* appears.",
            "Reopen the account menu — the name and the avatar initials both reflect the change.",
        ],
    )
)
A(("h3", "7.2.2  Phone number"))
A(
    (
        "p",
        "Optional. Enter a value and save, or clear the field and save — a blank phone is stored "
        "as empty, not rejected. Only the fields you actually changed are sent.",
    )
)
A(("h3", "7.2.3  Change your email"))
A(
    (
        "n",
        [
            "Edit **Email** to an unused address.",
            "Click **Save profile** — *Profile saved.*",
            "The account menu shows the new address, and **sign-in now requires it**. The old address no longer authenticates.",
        ],
    )
)
A(
    (
        "note",
        "Record the new address before you change it, especially on an administrator account. If "
        "you enter an address another user already holds you get `409` *That email is already in "
        "use.* and nothing changes.",
    )
)
A(("h3", "7.2.4  Change your password"))
A(
    (
        "p",
        "In the **Password** card enter your current password, the new password twice, and save. "
        "The failure messages are specific:",
    )
)
A(
    (
        "table",
        (
            ["Problem", "Message"],
            [
                ["Current password wrong", "*Current password is incorrect.* (`400`)"],
                ["New password under 12 characters", "*New password must be at least 12 characters.*"],
                ["Confirmation differs", "*New passwords do not match.*"],
                [
                    "Known-breached password",
                    "*password appears in a known-breach/common-password list; choose another* (`422`)",
                ],
                ["Success", "*Password changed.* and all three fields clear."],
            ],
            W2,
        ),
    )
)
A(("h2", "7.3  Multi-factor authentication"))
A(
    (
        "p",
        "MFA (TOTP) is fully implemented on the API — enrol, confirm, single-use recovery codes, "
        "disable, and an admin reset — but **there is no UI for any of it**, and the sign-in form "
        "has no code field. Enabling MFA on an account today makes that account unable to finish "
        "sign-in in the browser. See §26.7 for the endpoints and §28 for the gap.",
    )
)

# ── 8
sec("8  Finding your way around")
A(("h2", "8.1  The Dashboard"))
A(
    (
        "p",
        "`/` — reached by clicking the **DAS Sentinel** logo at the top of the sidebar or "
        "**Sidebar → Dashboard**. It shows the heading **Dashboard**, the line *Evidence-backed, "
        "compliance-mapped findings from authorized security testing.*, and a foundation card "
        "containing a **system health** link that opens `/health`.",
    )
)
A(("h2", "8.2  The sidebar"))
A(
    (
        "p",
        "Sections top to bottom, with the items each role can see:",
    )
)
A(
    (
        "table",
        (
            ["Section", "Item", "Visible to"],
            [
                ["Overview", "Dashboard", "Every role"],
                ["Testing", "Engagements", "Every role"],
                ["Credentials", "Credentials", "Admin, Tester"],
                ["Administration", "Users", "Admin"],
                ["Output", "Audit log", "Admin, Reviewer"],
                ["System", "AI models", "Every role (managing needs Admin)"],
                ["System", "Health", "Every role"],
            ],
            (1.6, 2.2, 2.8),
        ),
    )
)
A(
    (
        "p",
        "The sidebar is sticky — it stays visible while long pages scroll. Beneath the logo it "
        "reads *Authorized testing only*. There is no user block at the bottom of the sidebar; "
        "identity lives in the header avatar.",
    )
)
A(("h2", "8.3  The engagement context"))
A(
    (
        "p",
        "Targets, scans, findings and reports only exist inside an engagement, so they are not "
        "global sidebar entries. Open any engagement and a **Current engagement** section appears "
        "in the sidebar:",
    )
)
A(
    (
        "b",
        [
            "**Overview** — the engagement detail page.",
            "**Targets** and **Scans** — anchors that scroll the overview to those cards.",
            "**Findings** and **Reports** — their own pages.",
        ],
    )
)
A(
    (
        "p",
        "The section persists across the engagement's sub-pages and disappears when you leave the "
        "engagement. At the top of every engagement page you get the breadcrumb "
        "**Engagements / {name}** plus the status badge. Admin and Reviewer additionally see a "
        "**View audit log** link filtered to that engagement.",
    )
)

# ── 9
sec("9  Engagements")
A(
    (
        "p",
        "The engagement is the authorization container. Create it first; everything else hangs "
        "off it.",
    )
)
A(("h2", "9.1  The engagements list"))
A(
    (
        "p",
        "**Sidebar → Engagements**. Columns: **Name** (a link), **Client / system**, **Status** "
        "(badge), **Max intensity**, **Rate limit** (shown as *n rps*). Newest first. With none "
        "yet you get a dashed *No engagements yet…* box. The **New engagement** button sits "
        "top-right and is shown only to Admin and Tester.",
    )
)
A(("h2", "9.2  Create an engagement"))
A(
    (
        "n",
        [
            "Click **New engagement**.",
            "Fill in **Name** and **Client / system** — both required.",
            "Set **Test window start** and **Test window end**. See the warning below.",
            "Set **Rate limit (rps)** — between 1 and 1000, default 5.",
            "Choose **Maximum intensity** — the ceiling no scan in this engagement may exceed.",
            "Choose an **AI model** — a model registered under System → AI models, or *Organization default*.",
            "Tick **Hosted LLMs allowed** only if sending prompts off-box is authorized for this client.",
            "Fill in **Coordination contact** and **Emergency-stop contact**.",
            "Click **Create engagement**. The button reads *Saving…*, then you land on the new engagement's detail page in **Draft** status.",
        ],
    )
)
A(("h3", "9.2.1  Every field explained"))
A(
    (
        "table",
        (
            ["Field", "Required", "Notes"],
            [
                ["Name", "Yes", "Free text. Blocked in the browser if blank."],
                ["Client / system", "Yes", "The client or system under assessment. Blocked if blank."],
                [
                    "Test window start / end",
                    "No — but see the warning",
                    "End must be after start, or the save fails with *Some fields are invalid — check "
                    "the test window and rate limit.* (`422`).",
                ],
                [
                    "Rate limit (rps)",
                    "Yes",
                    "Default 5. Minimum 1, maximum 1000 — enforced in the browser. This is the "
                    "aggregate egress ceiling for the engagement.",
                ],
                [
                    "Maximum intensity",
                    "Yes",
                    "Exactly four options: **Passive**, **Safe active** (default), "
                    "**Authenticated active**, **High risk**. A launch above this is refused.",
                ],
                [
                    "AI model",
                    "No",
                    "Which registered model this engagement's triage, remediation and log analysis "
                    "run on. Blank = the organization default. With none registered the option reads "
                    "*No models registered — System → AI models*.",
                ],
                [
                    "Hosted LLMs allowed",
                    "No",
                    "**Unchecked by default** — local models only. Ticking it is what permits hosted "
                    "egress, and redaction still runs first. The detail page shows *Allowed* or "
                    "*Local models only*.",
                ],
                [
                    "Coordination contact",
                    "No",
                    "Who to call about the assessment. Free text, up to 500 characters.",
                ],
                [
                    "Emergency-stop contact",
                    "No",
                    "Who to call if testing must stop now. Free text, up to 500 characters.",
                ],
            ],
            (1.6, 1.2, 3.8),
        ),
    )
)
A(
    (
        "note",
        "**The test window is not optional in practice.** The form lets you leave it blank, but an "
        "absent window is treated exactly like an expired one: every scan is refused with "
        "`outside_test_window`. You can build a complete, correct-looking engagement that can "
        "never run a scan. Always set the window.",
    )
)
A(("h2", "9.3  The engagement detail page"))
A(("p", "One page carries the whole engagement. Cards, top to bottom:"))
A(
    (
        "b",
        [
            "**Details** — Client/system, both window timestamps (or *—*), Rate limit, Maximum intensity, Hosted LLMs, both contacts, Created and Updated timestamps.",
            "**Status** — the current badge plus the buttons for the transitions that are legal right now (§9.5).",
            "**Scope** — the allow and deny lists plus the add form (§10).",
            "**Rules of Engagement** — status badge, ROE text, content hash, and the acceptance control (§11).",
            "**Targets** — the inventory table plus **Add target** and **Add LLM target** (§12).",
            "**Scans** — the two launchers and *Recent scans* (§16).",
            "**Edit** and **Delete engagement** — shown only to Admin and Tester.",
        ],
    )
)
A(("h2", "9.4  Edit an engagement"))
A(
    (
        "n",
        [
            "Click **Edit** on the detail page.",
            "Change any field. Every create field is editable — **except status**, which is deliberately absent.",
            "Click **Save changes**. You return to the detail page with the new values.",
        ],
    )
)
A(
    (
        "note",
        "Editing a term that the ROE binds — the test window, the rate limit, or the maximum "
        "intensity — invalidates the acknowledgement. Scans are then refused with `roe_stale` or "
        "`roe_terms_mismatch` until you re-accept.",
    )
)
A(("h2", "9.5  Status transitions"))
A(
    (
        "p",
        "Status changes go through the Status card, never through the edit form. Only legal "
        "transitions are offered as buttons.",
    )
)
A(
    (
        "table",
        (
            ["Status", "What it means", "Available transitions"],
            [
                ["**Draft**", "Being set up. No scan can run.", "**Activate** → Active"],
                [
                    "**Active**",
                    "Testing is authorized, subject to the ROE, window and scope checks.",
                    "**Pause** → Paused · **Close** → Closed",
                ],
                ["**Paused**", "Temporarily halted. No scan can run.", "**Activate** → Active · **Close** → Closed"],
                [
                    "**Closed**",
                    "Terminal. The card reads *Closed is terminal — no further transitions.*",
                    "None. Forcing one over the API returns `409` and is audited.",
                ],
            ],
            (1.1, 2.9, 2.6),
        ),
    )
)
A(
    (
        "p",
        "**Close** asks you to confirm *Close this engagement? Closed is terminal.* Cancelling "
        "leaves the status untouched.",
    )
)
A(("h2", "9.6  Delete an engagement"))
A(
    (
        "n",
        [
            "On the detail page click **Delete engagement**.",
            "Confirm *Delete this engagement? It disappears from every list.*",
            "You return to `/engagements` and it is gone from the list.",
        ],
    )
)
A(
    (
        "p",
        "The delete is a soft delete — the row is retained for audit integrity, it simply stops "
        "appearing. Cancelling the dialog does nothing.",
    )
)
A(("h2", "9.7  What Reviewer and Read only see"))
A(
    (
        "p",
        "Both can open and read an engagement fully. Neither is offered **Edit**, the Status "
        "buttons, **Delete engagement**, or **New engagement**. Reaching `/engagements/{id}/edit` "
        "by URL and saving anyway returns `403` with *Your role can view engagements but not "
        "change them.* and changes nothing.",
    )
)

# ── 10
sec("10  Scope")
A(
    (
        "p",
        "Scope is the keystone safety control. Every target is matched against it immediately "
        "before execution — once when you launch, and again in the worker before the tool starts.",
    )
)
A(("h2", "10.1  The scope editor"))
A(
    (
        "p",
        "The **Scope** card on the engagement detail page holds two lists — **In scope (allow)** "
        "and **Out of scope (deny — always wins)** — each with an empty-state message, and an add "
        "form beneath them.",
    )
)
A(("h2", "10.2  Add a scope item"))
A(
    (
        "n",
        [
            "Choose **List**: *Allow* or *Deny*.",
            "Choose **Matcher type** (see §10.3).",
            "Enter the **Value**.",
            "Optionally add **Notes** — why this rule exists.",
            "Click **Add scope item**. The item appears in its list with its value, matcher label and notes.",
        ],
    )
)
A(
    (
        "p",
        "The Value and Notes fields clear after each add but List and Matcher stay selected, so "
        "entering a batch of rules of the same kind is fast.",
    )
)
A(("h2", "10.3  Matcher types"))
A(
    (
        "table",
        (
            ["Matcher", "Matches", "Valid example", "Rejected example"],
            [
                ["**URL**", "Scheme + host + path prefix", "`https://portal.example.com`", "`not-a-url`"],
                ["**Domain**", "The host and its subdomains", "`*.example.com`", "`bad_domain!`"],
                [
                    "**IP / CIDR**",
                    "An address inside the range — a hostname is resolved and its IP checked too",
                    "`10.0.0.0/24`",
                    "`999.1.1.1`",
                ],
                ["**API base**", "A URL prefix", "`https://api.x/v1`", "`ftp://x`"],
                ["**Repository**", "Normalized repository identity", "`git@github.com:o/r.git`", "free text"],
            ],
            (1.1, 2.3, 1.7, 1.5),
        ),
    )
)
A(
    (
        "p",
        "Valid values are accepted and normalized — lowercased, CIDRs canonicalized. An invalid "
        "value is refused with *The value is not a valid {type} matcher.*",
    )
)
A(
    (
        "note",
        "**Unresolvable or ambiguous always fails closed** — treated as out of scope. And an allow "
        "rule never authorizes a loopback, private, or cloud-metadata address: those are refused "
        "separately with `ssrf_ip_blocked`.",
    )
)
A(("h2", "10.4  Remove a scope item"))
A(("p", "Click **Remove** next to the item. It disappears from its list immediately."))
A(("h2", "10.5  Scope edits invalidate the ROE"))
A(
    (
        "p",
        "The ROE content hash covers the scope snapshot, so **any** scope change flips the ROE "
        "panel from *Accepted* back to *Acceptance required*, and scans are refused until you "
        "re-accept. This is deliberate: it guarantees every scan ran under a scope a human "
        "explicitly signed off.",
    )
)
A(("h2", "10.6  Role restrictions"))
A(
    (
        "p",
        "Only Admin and Tester can edit scope. Reviewer and Read only get *Your role can view "
        "scope but not change it.*",
    )
)

# ── 11
sec("11  Rules of Engagement")
A(("h2", "11.1  The ROE panel"))
A(
    (
        "p",
        "The **Rules of Engagement** card shows a status badge, the ROE text in a scrollable box, "
        "and a **Content hash: {hex}** line — a SHA-256 over the ROE text, the scope snapshot and "
        "the engagement terms.",
    )
)
A(("h2", "11.2  Accept the ROE"))
A(
    (
        "n",
        [
            "Read the ROE text. Actually read it — it is the authorization you are attesting to.",
            "Tick *I have read the Rules of Engagement…*. Until you do, the **Accept Rules of Engagement** button stays disabled.",
            "Click **Accept Rules of Engagement**.",
        ],
    )
)
A(
    (
        "p",
        "The badge turns green **Accepted** with the acceptance timestamp, and the checkbox and "
        "button disappear. Acceptance writes an immutable row — the database denies UPDATE and "
        "DELETE on it by trigger, not merely by hiding buttons.",
    )
)
A(("h2", "11.3  Re-acceptance"))
A(
    (
        "p",
        "An amber **Acceptance required** badge means the acknowledgement does not match the "
        "current terms. That happens when you edit scope, or change the test window, rate limit, "
        "or maximum intensity. Re-tick and re-accept; the badge returns to **Accepted** with a new "
        "timestamp, and a new acknowledgement row is written alongside the old one.",
    )
)
A(
    (
        "p",
        "The full history — every acknowledgement with its own content hash, acceptor and "
        "timestamp — is available on the API but has no screen (§26.5). It is what lets an auditor "
        "see exactly which ROE text and scope each scan ran under.",
    )
)
A(("h2", "11.4  Role restrictions"))
A(
    (
        "p",
        "Only Admin and Tester can accept. Reviewer and Read only get *Your role can view the ROE "
        "but not accept it.*",
    )
)

# ── 12
sec("12  Targets")
A(
    (
        "p",
        "A target is one system under test. Its **type** is what decides which scanners and "
        "suites can run against it, and the type cannot be changed later.",
    )
)
A(("h2", "12.1  The targets table"))
A(
    (
        "p",
        "On the engagement overview: **Name** (a link to its edit page, with the primary value "
        "beneath in monospace), **Type**, **Environment** (badge), **Auth**. Empty state: *No "
        "targets yet — add the systems this engagement is authorized to test.*",
    )
)
A(("h2", "12.2  Add a target"))
A(
    (
        "n",
        [
            "In the Targets card click **Add target** (or **Add LLM target**, which opens the same form).",
            "Enter a **Name**.",
            "Choose the **Type** — this drives the rest of the form.",
            "Fill the **primary value**; its label and placeholder change with the type (§12.4).",
            "Choose the **Environment** and **Auth status**.",
            "Optionally fill **Auth config** (§12.5) and, for LLM types, **Connector config** (§12.6).",
            "Click **Add target**. You return to the engagement and the target is listed.",
        ],
    )
)
A(("h2", "12.3  Target types"))
A(
    (
        "table",
        (
            ["Type", "Primary value", "What can test it"],
            [
                ["Web application (default)", "URL", "ZAP"],
                ["REST API", "URL", "ZAP"],
                ["GraphQL API", "URL", "ZAP"],
                ["Source repository", "Repository URL", "Semgrep"],
                ["Source archive", "Archive reference", "Semgrep — upload the archive first (§13)"],
                ["AI chatbot", "URL", "Prompt injection and Data leakage suites"],
                ["LLM API wrapper", "URL", "Prompt injection and Data leakage suites"],
                [
                    "AI agent",
                    "URL",
                    "The `agent_permission` suite — **API only**, no launcher accepts this type (§26.6)",
                ],
            ],
            (1.7, 1.5, 3.4),
        ),
    )
)
A(("h2", "12.4  Other fields"))
A(
    (
        "b",
        [
            "**Environment** — *Dev* (default), *Staging*, *Production*. Staging and Production get coloured badges; Production is red, because you should notice it.",
            "**Auth status** — *No auth* (default), *Configured*, *Verified*. Informational: it records how far you got setting up authenticated testing.",
        ],
    )
)
A(("h2", "12.5  Auth config — references only, never secrets"))
A(
    (
        "p",
        "**Auth config** is a JSON object. It must be a JSON *object* — an array or malformed JSON "
        "is blocked in the browser with *Auth config must be a JSON object (or empty).*",
    )
)
A(
    (
        "p",
        "Any `*_ref` value must use a `cred:`, `env:`, or `vault:` scheme. A raw secret is refused "
        "with `422`: `{\"api_key_ref\": \"hunter2\"}` fails, `{\"api_key_ref\": \"cred:<id>\"}` "
        "works.",
    )
)
A(("h3", "Using the credential picker"))
A(
    (
        "n",
        [
            "Create the secret first in the credentials vault (§14).",
            "On the target form open the **Credential** dropdown and pick it.",
            "The Auth config textarea gains `{\"api_key_ref\": \"cred:<id>\"}`, merged into whatever JSON is already there, and the picker resets.",
        ],
    )
)
A(("p", "The secret value itself is never inserted into the form and never displayed."))
A(("h2", "12.6  Connector config"))
A(
    (
        "p",
        "A **Connector config** textarea appears for **AI chatbot** and **LLM API wrapper** types "
        "only — it tells the platform how to talk to that model endpoint. It does **not** appear "
        "for AI agent. Invalid JSON is blocked in the browser; an unknown transport key is refused "
        "by the server with `422`.",
    )
)
A(("h2", "12.7  Edit a target"))
A(
    (
        "n",
        [
            "Click the target name in the table.",
            "Change the name, environment, auth status, auth config, or connector config.",
            "Click **Save changes**.",
        ],
    )
)
A(
    (
        "p",
        "**Type is immutable.** On the edit page the Type dropdown is disabled, with helper text "
        "explaining the type is fixed after creation. If the type is wrong, delete the target and "
        "add it again.",
    )
)
A(("h2", "12.8  Delete a target"))
A(
    (
        "p",
        "On the target's edit page click **Delete target** and confirm *Remove this target from "
        "the engagement's inventory?* It is soft-deleted and stops being listed.",
    )
)
A(("h2", "12.9  Role restrictions"))
A(
    (
        "p",
        "Only Admin and Tester can add, edit, or delete. Reviewer and Read only get *Your role can "
        "view targets but not change them.*",
    )
)

# ── 13
sec("13  Uploading a source archive")
A(
    (
        "p",
        "To run Semgrep against code you cannot reach by repository URL, create a **Source "
        "archive** target and upload the code to it.",
    )
)
A(("h2", "13.1  Steps"))
A(
    (
        "n",
        [
            "Add a target of type **Source archive**.",
            "Open its **Edit** page. A **Source archive** upload card appears — it shows only for this target type.",
            "Click the file chooser. It is filtered to `.zip`, `.tar`, `.tar.gz`, `.tgz`.",
            "Select the archive and click **Upload archive**. The button reads *Uploading…*",
            "On success a green line reports *Uploaded zip archive (N bytes) — sha256 <16 hex>…* and the target's primary value is set to the stored object key.",
        ],
    )
)
A(("h2", "13.2  What gets rejected"))
A(
    (
        "table",
        (
            ["Situation", "Result"],
            [
                ["No file chosen", "*Choose an archive (.zip or .tar) to upload.*"],
                ["Larger than 100 MiB", "`413` *That archive is too large.* Nothing is stored."],
                [
                    "Zip-slip, zip-bomb, symlink archive, or a non-archive renamed `.zip`",
                    "`422` *That file is not a valid, safe archive.* Nothing is stored.",
                ],
            ],
            (2.9, 3.7),
        ),
    )
)
A(
    (
        "p",
        "Archive validation is a hardened parser path, not a courtesy check — the platform ingests "
        "untrusted archives, so refusing hostile ones is a security control.",
    )
)

# ── 14
sec("14  Credentials vault")
A(
    (
        "p",
        "Where target secrets live. Visible to **Admin and Tester** only — Reviewer and Read only "
        "do not see the sidebar item and get `403` if they open `/credentials` directly.",
    )
)
A(("h2", "14.1  Create a credential"))
A(
    (
        "n",
        [
            "**Sidebar → Credentials**.",
            "Enter a **Name** (e.g. `prod-api-key`), an optional **Description**, and the **Secret**.",
            "Click **Create credential**.",
        ],
    )
)
A(
    (
        "p",
        "The credential appears in the list. **The secret is write-only** — it is encrypted at rest "
        "and never shown again, anywhere, to anyone. A duplicate name is refused with `409` *A "
        "credential with this name already exists.*",
    )
)
A(("h2", "14.2  Use a credential"))
A(
    (
        "p",
        "Each row shows a monospace reference `cred:<id>`. Click it to copy it to the clipboard — "
        "the row briefly shows *copied!*. Paste that reference into a target's **Auth config** as "
        "a `*_ref` value, or use the credential picker on the target form, which does it for you "
        "(§12.5).",
    )
)
A(("h2", "14.3  Delete a credential"))
A(
    (
        "p",
        "Click **Delete** on the row. It is soft-deleted. Only the name, description and "
        "`cred:<id>` are ever displayed in the list — there is no reveal control.",
    )
)

# ── 15
sec("15  AI models")
A(
    (
        "p",
        "**Sidebar → AI models**. Register a model once — a hosted provider API key, or a local "
        "Ollama endpoint — and engagements use it for triage, remediation and log analysis. Every "
        "role can view the page; **only an Admin can add, remove, or change the default**.",
    )
)
A(
    (
        "p",
        "That restriction is deliberate. Registering a model sets both a provider API key and an "
        "endpoint the platform itself calls out to — a tester who could point it anywhere would "
        "hold an authenticated egress primitive.",
    )
)
A(("h2", "15.1  Register a model"))
A(
    (
        "n",
        [
            "In **Add a model**, choose the **Provider**: *Anthropic Claude (hosted)* or *Ollama (local)*.",
            "Enter a **Name** — your label for it, e.g. `triage-model`.",
            "Enter the **Model** id exactly as the provider names it: `claude-opus-4-8` or `claude-sonnet-5` for Anthropic; whatever `ollama list` shows, e.g. `llama3.1:8b`, for Ollama.",
            "For Anthropic, paste the **API key**. For Ollama, give the **Ollama endpoint** (default `http://localhost:11434`).",
            "Click **Add model**.",
        ],
    )
)
A(
    (
        "p",
        "**The key and endpoint are validated against the provider before the model is saved** — a "
        "bad key or a model name the provider does not recognise is rejected at registration "
        "rather than failing later mid-analysis. The first model you register becomes the default "
        "automatically.",
    )
)
A(("h3", "Notes per provider"))
A(
    (
        "b",
        [
            "**Anthropic (hosted)** — the key is stored encrypted and is write-only. Prompts leave the box, so redaction runs before egress *and* each engagement must separately tick **Hosted LLMs allowed**.",
            "**Ollama (local)** — inference stays on-box. The model must already be pulled. For Ollama running on the same machine keep `localhost`; it resolves to the Docker host automatically.",
        ],
    )
)
A(("h2", "15.2  Registered models"))
A(("p", "Each row shows the name, the badges, and `provider · model_id · base_url` in monospace:"))
A(
    (
        "b",
        [
            "A green **default** badge on the model used when an engagement does not pick one.",
            "An amber **hosted · off-box** badge, or a plain **local · on-box** badge.",
        ],
    )
)
A(("h2", "15.3  Change the default, or remove a model"))
A(
    (
        "b",
        [
            "**Make default** — shown on every non-default row. Click it and that model becomes the organization default.",
            "**Remove** — deletes the registration. A model an engagement explicitly references is protected: the delete is refused with `409` rather than silently breaking that engagement.",
        ],
    )
)
A(("h2", "15.4  Errors you may see"))
A(
    (
        "table",
        (
            ["Code", "Message"],
            [
                ["`400`", "The provider itself rejected the key or the model name — shown verbatim."],
                ["`409`", "*A model with this name is already registered.*"],
                ["`403`", "*Only an admin can register AI models.*"],
            ],
            (0.9, 5.7),
        ),
    )
)
A(("h2", "15.5  When nothing is registered"))
A(
    (
        "p",
        "The list reads *No models registered yet.* and a card explains that analysis falls back "
        "to the deployment's environment configuration, showing that provider and default model. "
        "No API key or secret is ever displayed on this page.",
    )
)

# ── 16
sec("16  Running tests")
A(
    (
        "p",
        "Before any launch will succeed: the engagement is **Active**, the **ROE is accepted and "
        "current**, the target is **in scope**, **now is inside the test window**, and the "
        "intensity is **at or below the ceiling**. If a launch is refused, §17 decodes the reason.",
    )
)
A(("h2", "16.1  The Scans card"))
A(
    (
        "p",
        "Two launchers side by side — **AI / LLM suites** and **Code & web scanners** — above a "
        "**Recent scans** area that stays empty until something runs.",
    )
)
A(("h2", "16.2  Launch an AI / LLM suite"))
A(
    (
        "n",
        [
            "In the **AI / LLM suites** launcher pick a **Target**. The dropdown lists only AI chatbot and LLM API wrapper targets.",
            "Tick the **Suites** you want: **Prompt injection (LLM01)** — ticked by default — and **Data leakage (LLM02/05/07/08)**.",
            "Choose the **Intensity**: *Safe active* (default) or *Authenticated active*.",
            "Click **Launch scan**. The button reads *Launching…*, then a row appears in Recent scans as **Queued**, moving to **Running**.",
        ],
    )
)
A(("p", "Unticking every suite is refused with *Choose at least one suite to run.*"))
A(("h2", "16.3  Launch a code or web scanner"))
A(
    (
        "n",
        [
            "In **Code & web scanners** pick a **Target**.",
            "The **Scanners** checkboxes filter to the ones that apply to that target type, and the applicable ones are ticked by default: **Semgrep (SAST — source code)** for source repo and source archive targets; **ZAP (DAST — running web/API)** for web app, REST API and GraphQL targets.",
            "Choose the **Intensity**.",
            "Click **Launch scanner**.",
        ],
    )
)
A(
    (
        "p",
        "A mismatch is refused with `422` *{scanner} cannot run against a {type} target (supported: "
        "…)*. The UI only offers applicable scanners, so you normally cannot hit this from the "
        "browser.",
    )
)
A(("h3", "16.3.1  High-risk actions are not launchable here"))
A(
    (
        "p",
        "The scanner launcher carries a note that high-risk actions — exploit validation, brute "
        "force, destructive checks — require an approved high-risk gate and cannot be launched "
        "from this panel, with an **Approvals** link. **That link currently lands on the "
        "engagement overview, which has no approvals section**; there is no approvals UI anywhere. "
        "See §18 and §28.",
    )
)
A(("h2", "16.4  Watching a scan"))
A(
    (
        "p",
        "**Recent scans** columns: **Target**, **Intensity**, **Status**, **Queued** (time), and an "
        "action. Status badges are colour-coded: **Queued** (grey), **Running** (blue), "
        "**Completed** (green), **Failed** (red), **Cancelled** (amber).",
    )
)
A(
    (
        "p",
        "While any scan is active the panel shows *Live — updating while scans are active.* and "
        "refreshes about every 2.5 seconds on its own. No page refresh needed. Polling stops when "
        "nothing is active.",
    )
)
A(("h2", "16.5  Cancel a scan — the emergency stop"))
A(
    (
        "n",
        [
            "Click **Cancel** on the running scan's row.",
            "The button reads *Stopping…* and a hint appears.",
            "The scan moves to **Cancelled**.",
        ],
    )
)
A(
    (
        "p",
        "Cancellation is not cosmetic. The worker signals the recorded process group `SIGTERM` then "
        "`SIGKILL`, **confirms the process tree is actually gone**, verifies the sandbox was torn "
        "down, marks the scan cancelled, and audits it. An in-process AI suite honours a "
        "cooperative cancellation token. A scan is cancellable while **Queued** or **Running**.",
    )
)
A(
    (
        "note",
        "If a run appears to reach anything outside the sandbox or the engagement's allowlist: "
        "**cancel it, then escalate before continuing.** That is what the emergency-stop contact "
        "on the engagement is for.",
    )
)
A(("h2", "16.6  Limits you will run into"))
A(
    (
        "table",
        (
            ["Limit", "Value", "What happens"],
            [
                [
                    "Concurrent scans per engagement",
                    "5",
                    "Excess launches refused with `429` and `Retry-After: 30`. Already-queued scans still run.",
                ],
                ["Concurrent scans per organization", "20", "Same refusal."],
                [
                    "Egress rate",
                    "The engagement's **Rate limit (rps)**",
                    "Outbound traffic is shaped in aggregate at the engagement's egress choke point.",
                ],
                [
                    "State-changing API requests per user",
                    "300 per minute",
                    "`429` with `Retry-After`. Reads are exempt.",
                ],
                [
                    "LLM spend per engagement",
                    "A configured token/cost ceiling",
                    "Once spent, further LLM calls are refused with `429` rather than silently running up cost.",
                ],
            ],
            (2.1, 1.5, 3.0),
        ),
    )
)
A(("h2", "16.7  Rules about what a scan may contain"))
A(
    (
        "b",
        [
            "**Exactly one of suites or scanners.** Sending both, or neither, is refused with `422` *provide exactly one of 'suites' or 'scanners'*. An AI-suite scan and a scanner scan are never mixed into one record.",
            "**`agent_permission` runs alone.** Combining it with a PyRIT suite is refused with `422` *agent_permission must be launched on its own* — it runs a different engine.",
            "**Only Semgrep and ZAP are launchable over HTTP.** Nuclei, OSV-Scanner, Gitleaks, testssl, httpx and katana adapters exist in the worker but are not exposed by the launch API; a request naming one is refused as an invalid value. See §28.",
        ],
    )
)
A(("h2", "16.8  Role restrictions"))
A(
    (
        "p",
        "Only Admin and Tester can launch or cancel. Reviewer and Read only are refused with "
        "`403`.",
    )
)

# ── 17
sec("17  Why a launch was blocked")
A(
    (
        "p",
        "Every refusal carries a stable machine reason and is written to the audit log with "
        "outcome `blocked`. This is the table to read when a scan will not start.",
    )
)
A(
    (
        "table",
        (
            ["Reason", "HTTP", "What it means", "How to fix it"],
            BLOCK_REASONS,
            (1.5, 0.5, 2.3, 2.3),
        ),
    )
)
A(("h2", "17.1  Two more refusals worth knowing"))
A(
    (
        "b",
        [
            "**Hosted model when the engagement forbids it** — `409`. Only local models may be used for that engagement. Tick **Hosted LLMs allowed** on the engagement, if that is genuinely authorized.",
            "**Redaction failed before hosted egress** — the call is blocked, not retried unredacted. Fail-closed is the intended behaviour.",
        ],
    )
)
A(("h2", "17.2  The check runs twice"))
A(
    (
        "p",
        "The same authorization function runs in the API before the job is enqueued, and again in "
        "the worker before the tool launches — the second time re-reading the database and the "
        "immutable authorization envelope and re-deriving everything, trusting nothing it was "
        "handed. A window that closes between enqueue and execution stops the run.",
    )
)
A(("h2", "17.3  Cross-engagement isolation"))
A(
    (
        "p",
        "Changing an id in a URL to an engagement, target, scan, finding, report, approval or "
        "evidence object you should not see returns **`404`** — not `403`, and never the data. "
        "You cannot use error codes to enumerate what exists.",
    )
)

# ── 18
sec("18  High-risk approval gates")
A(
    (
        "note",
        "**There is no approvals screen.** Everything in this section is done over the API, and at "
        "this milestone **no HTTP path spends an approved gate** — consumption happens inside the "
        "worker. An approved gate therefore does not yet change what you can launch from the UI or "
        "the API. Documented here because the lifecycle is fully implemented and audited.",
    )
)
A(("h2", "18.1  What needs a gate"))
A(
    (
        "p",
        "Four operation kinds derive high-risk intensity and therefore need one: "
        "`exploit_validation`, `brute_force`, `large_crawl`, `data_modifying`. Requesting a gate "
        "for anything else — e.g. `safe_active_scan` — is refused with `400` *approval applies "
        "only to high-risk operations*. **Intensity is derived server-side from the operation "
        "kind; a caller never declares it.**",
    )
)
A(("h2", "18.2  Request a gate"))
A(
    (
        "code",
        'POST /api/engagements/{engagement_id}/approvals\n'
        '{\n'
        '  "target_id": "<target uuid>",\n'
        '  "operation_kind": "exploit_validation",\n'
        '  "justification": "why this is necessary",\n'
        '  "expires_in_hours": 24\n'
        '}',
    )
)
A(
    (
        "p",
        "`201` returns the gate with status `pending`, a hex `operation_digest`, the `roe_ack_id` "
        "bound to the **current** ROE acknowledgement, the policy version, and `expires_at`. "
        "Minimum expiry is 1 hour.",
    )
)
A(
    (
        "p",
        "Without a current accepted ROE the request is refused with `409` *a current accepted ROE "
        "is required before requesting an approval*. A stale acknowledgement cannot authorize a "
        "high-risk action.",
    )
)
A(("h2", "18.3  Decide, deny, revoke"))
A(
    (
        "table",
        (
            ["Action", "Endpoint and body", "Result"],
            [
                [
                    "Approve",
                    "`POST …/approvals/{id}/decide` `{\"approve\": true, \"reason\": \"…\"}`",
                    "Status `approved`, with `decided_by`, `decided_at`, `decision_reason`.",
                ],
                [
                    "Deny",
                    "`POST …/approvals/{id}/decide` `{\"approve\": false, \"reason\": \"…\"}`",
                    "Status `denied`.",
                ],
                [
                    "Decide twice",
                    "the same call again",
                    "`409` *cannot decide an approval in state approved* — decisions are one-shot.",
                ],
                [
                    "Revoke",
                    "`POST …/approvals/{id}/revoke` `{\"reason\": \"…\"}`",
                    "An `approved` gate becomes `revoked`. Revoking anything else → `409`.",
                ],
            ],
            (0.9, 2.9, 2.8),
        ),
    )
)
A(("h2", "18.4  Expiry and single use"))
A(
    (
        "b",
        [
            "A gate past its `expires_at` auto-transitions to `expired` on the next touch, and deciding it is refused with *approval request has expired*. A periodic job also sweeps expiries in bulk.",
            "Consumption is **single-use and atomic**: of two concurrent attempts to spend the same gate, exactly one wins. An approval cannot be replayed.",
        ],
    )
)
A(("h2", "18.5  Who can do what"))
A(
    (
        "b",
        [
            "**Requesting** needs *launch scans* → **Admin and Tester** only.",
            "**Deciding and revoking** need *approve high-risk* → **Admin and Reviewer** only.",
            "**Viewing** needs *view* → all four roles.",
        ],
    )
)
A(
    (
        "note",
        "An Admin holds both capabilities, so **an Admin can approve their own request** — there "
        "is no separation-of-duties check that the approver differs from the requester. If you "
        "need four-eyes on high-risk authorization, enforce it by having Testers request and "
        "Reviewers approve, as a matter of policy. It is not currently a code control.",
    )
)
A(("h2", "18.6  Audit trail"))
A(
    (
        "p",
        "The audit log carries `approval.requested`, `approval.approved` / `approval.denied`, and "
        "`approval.revoked`, each with the actor, the target id, and the reason in the detail. "
        "Fetching a gate that belongs to another engagement returns `404`.",
    )
)

# ── 19
sec("19  Findings")
A(("h2", "19.1  The findings list"))
A(
    (
        "p",
        "**Current engagement → Findings**. The header shows *N finding(s)*. Empty state: *No "
        "findings yet — run an AI security scan against an in-scope target.* Intro copy reminds "
        "you that automated and AI findings are labelled and are not human-validated.",
    )
)
A(
    (
        "p",
        "Columns: **Severity**, **Finding**, **Source**, **OWASP**, **Provenance**, **Status**. "
        "Ordering is severity-first — Critical down to Info — and newest first within each "
        "severity band.",
    )
)
A(("h2", "19.2  Reading the badges"))
A(("h3", "Severity"))
A(("p", "**Critical** · **High** · **Medium** · **Low** · **Info**, each with its own colour."))
A(("h3", "Provenance — the truthfulness control"))
A(
    (
        "table",
        (
            ["Badge", "Meaning"],
            [
                ["**Automated**", "Produced by a deterministic detector — not human-validated."],
                ["**AI-generated**", "Draft AI analysis — not human-validated."],
                ["**Validated**", "Reviewed and confirmed by a human."],
                ["**Manually overridden**", "A human changed this finding from its original state."],
            ],
            (1.9, 4.7),
        ),
    )
)
A(
    (
        "p",
        "An AI-generated finding awaiting review also carries an amber **review** tag, tooltip "
        "*AI-proposed — needs human review*. The **OWASP** column shows the reference code with "
        "the framework and title as a tooltip.",
    )
)
A(("h3", "Status"))
A(
    (
        "p",
        "**Open** · **In triage** · **Confirmed** · **Mitigated** · **Fixed** · **Accepted risk** "
        "· **False positive** · **Out of scope**.",
    )
)
A(
    (
        "note",
        "**These statuses are display-only at this milestone.** There is no UI or API path to "
        "transition a finding's status, so every finding stays as created and status history holds "
        "exactly one entry. The triage and review workflow the Reviewer role implies is not yet "
        "reachable. See §28.",
    )
)
A(("h2", "19.3  The finding detail page"))
A(("p", "Click a finding title. The page shows, in order:"))
A(
    (
        "b",
        [
            "**Header** — Severity, Provenance and Status badges.",
            "**Summary** — the message, Rule, OWASP LLM reference, Technique, Suite, Created timestamp, and optional Description and Recommendation.",
            "**An unvalidated notice** on any automated or AI-generated finding: it has not been human-validated, review the evidence before acting.",
            "**Human review required** — on an AI-generated finding, a checklist: verify the quoted evidence, confirm true or false positive, assign a real CVSS severity.",
            "**CVSS** — §20.",
            "**Compliance mappings** — §21.",
            "**Evidence** — §22.",
            "**Status history** — append-only; currently the single creation entry.",
        ],
    )
)

# ── 20
sec("20  CVSS scoring")
A(
    (
        "p",
        "CVSS v4.0 is the default for new findings; v3.1 is retained for historical CVEs. Scores "
        "are computed **server-side** from the vector you supply — the AI never sets a final CVSS.",
    )
)
A(("h2", "20.1  Reading the CVSS card"))
A(
    (
        "p",
        "When scored it shows the base score to one decimal, a severity-band badge, the version "
        "label (**CVSS v4.0** or **CVSS v3.1**), the vector string, and — if overridden — "
        "*· manual override* plus the justification. When unscored it reads *Not scored yet.* with "
        "an edit form, if your role can validate.",
    )
)
A(("h2", "20.2  Score a finding"))
A(
    (
        "n",
        [
            "Open the finding and find the **CVSS** card.",
            "Paste a full vector string, e.g. `CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:H/VA:H/SC:H/SI:H/SA:H`.",
            "Click **Save score**.",
        ],
    )
)
A(
    (
        "p",
        "The base score, severity band and version are derived server-side and displayed. A "
        "malformed vector is refused with `422` *That is not a valid CVSS vector.* and nothing is "
        "written.",
    )
)
A(("h2", "20.3  Manual override"))
A(
    (
        "p",
        "Tick **Manual override** to record a human judgement that differs from the computed "
        "score. A justification is mandatory — saving without one is blocked with *A manual "
        "override needs a justification.* Once saved, the card shows *· manual override* alongside "
        "your reasoning.",
    )
)
A(("h2", "20.4  Score history"))
A(
    (
        "p",
        "Re-scoring never overwrites. Expand the score-history disclosure to see every entry — "
        "base score, version, *(override)* marker, timestamp — insert-only.",
    )
)
A(("h2", "20.5  Role restrictions"))
A(
    (
        "p",
        "Admin, Tester and Reviewer can score. Read only sees no edit form and is refused.",
    )
)

# ── 21
sec("21  Compliance mapping")
A(
    (
        "p",
        "Mapping a finding to published controls is what makes the output defensible. Seven "
        "frameworks ship with the platform:",
    )
)
A(
    (
        "table",
        (
            ["Framework", "Version", "Controls", "Codes"],
            FRAMEWORKS,
            (2.5, 1.5, 0.7, 1.9),
        ),
    )
)
A(
    (
        "p",
        "If the Framework dropdown is empty, the knowledge base was never seeded — run "
        "`scripts/seed_compliance.py` (§5.2).",
    )
)
A(("h2", "21.1  Reading the card"))
A(
    (
        "p",
        "Mapped controls appear as tags showing the code and framework, with the framework name "
        "and control title as a tooltip. With none: *No control mappings yet.*",
    )
)
A(("h2", "21.2  Auto-map from a finding"))
A(
    (
        "p",
        "Click **Auto-map from finding**. Mappings are derived from the finding's structured "
        "references, and each carries an **auto** marker. There is also an engagement-wide "
        "auto-map on the API that does this across every finding at once (§26.4).",
    )
)
A(("h2", "21.3  Add a mapping by hand"))
A(
    (
        "n",
        [
            "Pick a **Framework**.",
            "Pick a **Control** — already-mapped controls are disabled in the picker.",
            "Click **Add mapping**.",
        ],
    )
)
A(
    (
        "p",
        "A hand-added mapping has no **auto** marker: that is how a reviewer tells an automatic "
        "inference from a human judgement.",
    )
)
A(("h2", "21.4  Remove a mapping"))
A(("p", "Click the **×** on the mapping tag. The card refreshes."))
A(("h2", "21.5  Role restrictions"))
A(
    (
        "p",
        "Admin, Tester and Reviewer can map. Read only sees no Auto-map, Add, or Remove controls.",
    )
)

# ── 22
sec("22  Evidence and transcripts")
A(
    (
        "p",
        "Evidence is the chain of custody. Raw artefacts live in the object store, content-hashed, "
        "under object-lock/WORM — never in a database column.",
    )
)
A(("h2", "22.1  The evidence list"))
A(
    (
        "p",
        "Each row on a finding shows a caption plus `kind · N bytes · sha256: <16 hex>…` and a "
        "**View transcript** toggle. With none: *No evidence attached.* Kinds are raw scanner "
        "output, HTTP transcript, LLM transcript, and source archive.",
    )
)
A(("h2", "22.2  View a transcript"))
A(
    (
        "n",
        [
            "Click **View transcript**. It loads lazily.",
            "An LLM transcript renders as conversation turns — a role header and its content.",
            "Click **Show raw JSON** to see the pretty-printed source, **Show conversation** to go back.",
        ],
    )
)
A(
    (
        "p",
        "Non-transcript evidence — a scanner's JSON output, say — falls back to the pretty-printed "
        "raw view.",
    )
)
A(("h2", "22.3  Integrity is re-checked on every read"))
A(
    (
        "p",
        "The panel notes *Verified via the evidence store (SHA-256 checked on read).* That is "
        "literal: the stored hash is recomputed before content is returned, and a mismatch fails "
        "closed rather than serving unverified bytes. Evidence is never mutated or deleted — the "
        "production database role is denied UPDATE and DELETE on the append-only tables.",
    )
)

# ── 23
sec("23  Reports and exports")
A(("h2", "23.1  The reports page"))
A(
    (
        "p",
        "**Current engagement → Reports**. A **Generate a report** card sits above a table of "
        "**Title**, **Type**, **Status**, **Created**. Empty: *No reports yet.*",
    )
)
A(("h2", "23.2  Generate a report"))
A(
    (
        "n",
        [
            "Choose the **Type**: **POA&M** (default), **Technical report**, or **Executive summary**.",
            "Enter a **Title**. Leaving it blank is refused with *Enter a report title.*",
            "Click **Generate report**. The button reads *Generating…*, then you land on the report builder with the report in **Draft**.",
        ],
    )
)
A(
    (
        "p",
        "Generating snapshots the engagement's findings into the report, so a later scan does not "
        "silently rewrite a report you already circulated.",
    )
)
A(("h2", "23.3  Edit a draft"))
A(
    (
        "b",
        [
            "**Summary** — type into the textarea and click **Save changes**; you get *Saved.* Editable only while the report is Draft.",
            "**POA&M per-finding fields** — for each finding, a **Responsible owner**, a **Planned completion date** (date picker), **Milestones**, and **Risk acceptance notes**. Fill them in and save.",
        ],
    )
)
A(("h2", "23.4  Finalize"))
A(
    (
        "p",
        "Click **Finalize**. The report locks: a banner says it is finalized and read-only, the "
        "edit fields disappear, and only **Delete** and the downloads remain. **Exports still "
        "work after finalization** — the banner says so.",
    )
)
A(("h2", "23.5  Export"))
A(("p", "Five formats, each a button on the builder. All work on a Draft or a Final report:"))
A(
    (
        "table",
        (
            ["Button", "File", "Content"],
            [
                [
                    "**Download POA&M CSV**",
                    "`poam-<id>.csv`",
                    "Plan of Action & Milestones: weakness id, asset, severity, CVSS, control mapping, owner, dates, status, milestones.",
                ],
                [
                    "**Download Markdown**",
                    "`report-<id>.md`",
                    "The technical report — findings, evidence, reproduction, remediation — or the executive narrative if the report type is Executive.",
                ],
                ["**Download PDF**", "`report-<id>.pdf`", "The same content rendered to PDF."],
                ["**Download DOCX**", "`report-<id>.docx`", "The same content as a Word document."],
                [
                    "**Download JSON**",
                    "`report-<id>.json`",
                    "The raw editable report body — for feeding another system.",
                ],
            ],
            (1.9, 1.5, 3.2),
        ),
    )
)
A(
    (
        "p",
        "Rendering is a pure function of the stored body, so the same report always exports "
        "identically. Every export writes a `report.exported` audit event recording the format.",
    )
)
A(("h2", "23.6  Delete a report"))
A(("p", "Click **Delete** and confirm *Delete this report? This cannot be undone.*"))
A(("h2", "23.7  Role restrictions"))
A(
    (
        "p",
        "Admin, Tester and Reviewer can generate, edit, finalize, export and delete. Read only "
        "sees *You do not have permission to generate reports.* and can view but not export.",
    )
)

# ── 24
sec("24  Audit log")
A(
    (
        "p",
        "**Sidebar → Audit log**, visible to **Admin and Reviewer**. Tester and Read only get *The "
        "audit log is an oversight view available to Admin and Reviewer roles only.*",
    )
)
A(("h2", "24.1  Reading the log"))
A(
    (
        "p",
        "Columns: **Time**, **Actor** (email, or *system*), **Action** with the object type and "
        "the detail JSON as a tooltip, **Engagement** (a link, or *—*), **Outcome** badge, and "
        "**IP**. Newest first.",
    )
)
A(
    (
        "b",
            [
            "**success** — outline badge.",
            "**blocked** — red. Every refused scan, every scope violation.",
            "**failure** — amber.",
        ],
    )
)
A(("h2", "24.2  Filter to one engagement"))
A(
    (
        "p",
        "Click an engagement name in the Engagement column. The header shows *Filtered to "
        "engagement <name> — [show all]*; **show all** clears it. Admin and Reviewer also get a "
        "**View audit log** link on each engagement page that arrives pre-filtered.",
    )
)
A(("h2", "24.3  What gets audited"))
A(
    (
        "table",
        (
            ["Area", "Events"],
            [
                ["Authentication", "`auth.login`, `auth.logout`, `session.revoked`"],
                [
                    "Account",
                    "`auth.profile_updated`, `auth.password_changed`",
                ],
                [
                    "Engagement, scope, ROE",
                    "`engagement.created` / `.updated`, `scope.updated`, `roe.accepted`",
                ],
                [
                    "Safety refusals",
                    "`scope.blocked` with the machine reason, outcome `blocked`",
                ],
                ["Approvals", "`approval.requested`, `.approved` / `.denied`, `.revoked`"],
                ["Scans", "`scan.queued`, `.started`, `.completed` / `.failed` / `.cancelled`"],
                [
                    "Findings",
                    "`finding.validated`, `.status_changed`, `cvss.overridden`, `finding.false_positive`",
                ],
                ["LLM", "`llm.call`, with the hosted and redacted flags"],
                ["Reports", "`report.generated`, `report.exported`"],
            ],
            (1.9, 4.7),
        ),
    )
)
A(("h2", "24.4  Immutability"))
A(
    (
        "p",
        "There is no edit or delete control on any audit row, and that is not just UI. UPDATE and "
        "DELETE are denied on `audit_events` and `roe_acknowledgements` by a database trigger — "
        "append-only is enforced below the application. Aged events are archived to object storage "
        "under object-lock so retention does not depend on the primary database.",
    )
)
A(
    (
        "note",
        "**The page shows only the newest 100 events and has no paging control.** Older events are "
        "reachable only through the API (`GET /api/audit-events?limit=…&offset=…`, limit capped at "
        "500). Worth knowing before you rely on this screen for a compliance review.",
    )
)

# ── 25
sec("25  System health")
A(
    (
        "p",
        "**Sidebar → Health**, or the **system health** link on the Dashboard. Available to every "
        "role.",
    )
)
A(
    (
        "b",
        [
            "Probes in order: **API** (liveness), **Database**, **Valkey** — each with an **ok** badge when healthy.",
            "A header badge shows **ok** (green) when every probe passes, **unavailable** (red) otherwise.",
            "Probes re-run on page load. There is no auto-refresh — reload to re-check.",
            "Stop a backing service and its probe turns red and the header badge flips.",
        ],
    )
)
A(
    (
        "note",
        "**The evidence store is not health-probed.** Readiness covers the database and Valkey "
        "only, so an object-storage outage leaves this page showing **ok** while evidence uploads "
        "and reads fail. If evidence behaves oddly, check object storage directly.",
    )
)

# ── 26
sec("26  Features that are API-only")
A(
    (
        "p",
        "These are implemented, tested and audited, but have no screen. Authenticate the way the "
        "browser does — the session cookie — and send `X-CSRF-Token` on every POST, PATCH and "
        "DELETE. The token is returned in the login response body and set as a cookie. A cookie "
        "without the header is refused by design.",
    )
)
A(("h2", "26.1  Log analysis"))
A(
    (
        "code",
        "POST /api/engagements/{eid}/targets/{tid}/log-analysis\n{\"evidence_id\": \"<raw_scanner_output evidence id>\"}",
    )
)
A(
    (
        "p",
        "Turns raw scanner output into candidate findings, created as **AI-generated / "
        "Informational / Open** — never presented as verified. Needs *validate findings* (Admin, "
        "Tester, Reviewer). Wrong evidence kind → `422`; missing or cross-organization → `404`; "
        "hosted model not allowed → `409`; budget exhausted → `429`.",
    )
)
A(("h2", "26.2  Remediation generation"))
A(
    (
        "code",
        "POST /api/engagements/{eid}/findings/{fid}/remediation/generate\nGET  /api/engagements/{eid}/findings/{fid}/remediation",
    )
)
A(
    (
        "p",
        "Creates an AI-generated **draft** remediation — a secure code example or patch suggestion "
        "— for human review. **The finding itself is never mutated.** Note the split: generating "
        "needs *validate findings*, but **listing is view-gated**, so Read only can read AI "
        "drafts.",
    )
)
A(("h2", "26.3  SARIF export and import"))
A(
    (
        "b",
        [
            "`GET /api/engagements/{eid}/findings/export-sarif` — a SARIF 2.1.0 log: version `2.1.0`, the standard `$schema`, one run, results carrying `ruleId`, `level` (error/warning/note) and locations. Canonical findings only, no duplicates.",
            "`POST /api/engagements/{eid}/findings/import-sarif` — creates automated findings with evidence. **Re-importing the same file deduplicates** rather than re-inserting. Malformed → `422`; oversized → `413`. Needs *manage engagements*.",
        ],
    )
)
A(
    (
        "note",
        "The SARIF **export** route is only view-gated, so a Read-only account can pull every "
        "finding as SARIF even though it is refused report export. Worth a policy decision if that "
        "asymmetry matters to you.",
    )
)
A(("h2", "26.4  Triage overview, scan plan, engagement-wide auto-map"))
A(
    (
        "b",
        [
            "`GET /api/engagements/{eid}/triage` — the deterministic triage view: findings ranked and grouped. Same input always yields the same order; **the AI does not decide ranking**.",
            "`GET /api/engagements/{eid}/targets/{tid}/scan-plan` — a deterministic plan derived from the target and its existing findings. No LLM in the ranking path. View-gated.",
            "`POST /api/engagements/{eid}/compliance/auto-map` — auto-map every finding in the engagement in one call.",
        ],
    )
)
A(("h2", "26.5  ROE acknowledgement history"))
A(
    (
        "p",
        "`GET /api/engagements/{eid}/roe/acknowledgements` returns every acknowledgement in order, "
        "each with its own content hash, acceptor and timestamp — so an auditor can see exactly "
        "which ROE text and scope each scan ran under. Insert-only. No screen shows it.",
    )
)
A(("h2", "26.6  The agent-permission suite"))
A(
    (
        "p",
        "The third shipped test suite drives an **AI agent** target through sandboxed fake tools "
        "and monitors its tool calls, producing permission-boundary findings with transcript "
        "evidence. You can create an `ai_agent` target in the UI, but no launcher accepts it:",
    )
)
A(
    (
        "code",
        'POST /api/engagements/{eid}/scans\n'
        '{"target_id": "<ai_agent target>", "suites": ["agent_permission"], "intensity": "safe_active"}',
    )
)
A(
    (
        "p",
        "Every keystone gate still applies. It must be launched on its own — combining it with a "
        "PyRIT suite is refused with `422`. **Every run drives the fake tools in `sandbox/`** — no "
        "real credential, mailbox, filesystem or third-party API is touched.",
    )
)
A(("h2", "26.7  MFA and admin user operations"))
A(
    (
        "table",
        (
            ["Endpoint", "What it does"],
            [
                [
                    "`POST /api/auth/mfa/enroll`",
                    "Stores a **pending** TOTP secret; `mfa_enabled` stays false until confirmed.",
                ],
                [
                    "`POST /api/auth/mfa/confirm`",
                    "Needs a live TOTP code; issues single-use recovery codes. Sign-in then demands the second factor.",
                ],
                ["`POST /api/auth/mfa/disable`", "Requires a valid factor."],
                [
                    "`POST /api/users/{id}/reset-mfa`",
                    "Admin-only. Clears MFA and every recovery code so a lost authenticator can be recovered, and revokes all that user's sessions. Audited.",
                ],
                [
                    "`POST /api/users/{id}/reset-password`",
                    "Admin-only. Returns a new one-time temporary password, re-arms the forced change, and **revokes every session that user holds** — an active session is kicked out immediately.",
                ],
            ],
            (2.3, 4.3),
        ),
    )
)
A(("h2", "26.8  Self-service account endpoints"))
A(
    (
        "p",
        "`GET /api/auth/me` returns the current principal — id, email, display name, role, phone, "
        "`must_change_password`. `PATCH /api/auth/me` touches only the fields you send and returns "
        "`409` on a duplicate email. `POST /api/auth/me/password` needs the current password unless "
        "the account is in forced-change mode. All three are audited.",
    )
)

# ── 27
sec("27  Complete API reference")
A(
    (
        "p",
        "Every route the platform exposes, with the capability it requires. Paths are relative to "
        "`/api`. A dash in the Capability column means the route is not behind a capability guard: "
        "`POST /auth/login` is unauthenticated by necessity, and the logout, self-service account, "
        "and MFA routes act on the caller's **own** account, which any signed-in user may do. "
        "`/healthz` and `/readyz` are unauthenticated.",
    )
)
A(
    (
        "p",
        "Read the Capability column against the matrix in §4.2: `VIEW` = all four roles; "
        "`MANAGE_ENGAGEMENTS`, `LAUNCH_SCANS`, `ACCEPT_ROE`, `MANAGE_CREDENTIALS` = Admin + "
        "Tester; `APPROVE_HIGH_RISK` = Admin + Reviewer; `VALIDATE_FINDINGS`, `EXPORT_REPORTS` = "
        "Admin + Tester + Reviewer; `VIEW_AUDIT` = Admin + Reviewer; `MANAGE_USERS`, "
        "`MANAGE_AI_MODELS` = Admin only.",
    )
)

GROUPS = [
    ("Health", ["GET /healthz", "GET /readyz"]),
    ("Authentication and account", None),
    ("Users", None),
    ("Engagements", None),
    ("Scope", None),
    ("Rules of Engagement", None),
    ("Targets and source archive", None),
    ("Credentials", None),
    ("Approvals", None),
    ("Scans", None),
    ("Findings and evidence", None),
    ("CVSS", None),
    ("Compliance", None),
    ("Remediation and log analysis", None),
    ("Triage and scan plan", None),
    ("Reports", None),
    ("AI models and LLM", None),
    ("Audit", None),
]

PREFIX_GROUP = [
    ("/auth", "Authentication and account"),
    ("/users", "Users"),
    ("/credentials", "Credentials"),
    ("/llm", "AI models and LLM"),
    ("/audit-events", "Audit"),
    ("/compliance/frameworks", "Compliance"),
]


def group_for(route: str) -> str:
    method, path = route.split(" ", 1)
    for pre, name in PREFIX_GROUP:
        if path.startswith(pre):
            return name
    if "/scope-items" in path:
        return "Scope"
    if "/roe" in path:
        return "Rules of Engagement"
    if "/approvals" in path:
        return "Approvals"
    if "/scans" in path:
        return "Scans"
    if "/compliance" in path:
        return "Compliance"
    if "/cvss" in path:
        return "CVSS"
    if "/remediation" in path or "/log-analysis" in path:
        return "Remediation and log analysis"
    if "/triage" in path or "/scan-plan" in path:
        return "Triage and scan plan"
    if "/reports" in path:
        return "Reports"
    if "/findings" in path:
        return "Findings and evidence"
    if "/targets" in path or "/source-archive" in path:
        return "Targets and source archive"
    return "Engagements"


ORDER = [g for g, _ in GROUPS]
grouped: dict[str, list[tuple[str, str]]] = {g: [] for g in ORDER}
grouped["Health"] = [("GET /healthz", "unauthenticated"), ("GET /readyz", "unauthenticated")]
for route, cap in CAPS.items():
    grouped[group_for(route)].append((route, cap))

for g in ORDER:
    rows = grouped[g]
    if not rows:
        continue
    A(("h3", g))
    A(
        (
            "table",
            (
                ["Method and path", "Capability"],
                [
                    [f"`{r.split(' ', 1)[0]}` `{r.split(' ', 1)[1]}`", f"`{c}`" if c != "-" else "—"]
                    for r, c in sorted(rows, key=lambda x: (x[0].split(" ", 1)[1], x[0]))
                ],
                (4.6, 2.0),
            ),
        )
    )

# ── 28
sec("28  Current limitations")
A(
    (
        "p",
        "Correct-as-built but incomplete. None of these is a malfunction; all of them change how "
        "you work today.",
    )
)
A(
    (
        "table",
            (
            ["Area", "Limitation", "Work around it by"],
            [
                [
                    "Finding status",
                    "No API or UI transitions a finding's status. The eight statuses are "
                    "display-only and status history holds one entry. The triage/review workflow "
                    "the Reviewer role implies is not reachable.",
                    "Recording review decisions outside the platform for now. CVSS scoring and "
                    "compliance mapping do work and are the durable record.",
                ],
                [
                    "Approvals",
                    "No approvals UI anywhere. The scanner launcher's **Approvals** link lands on "
                    "the engagement overview. No HTTP path spends an approved gate, so high-risk "
                    "stays blocked even after approval.",
                    "Driving the lifecycle over the API (§18) for the audit record.",
                ],
                [
                    "MFA",
                    "No UI at all, and the sign-in form has no code field — an MFA-enabled account "
                    "cannot finish sign-in in the browser, and shows only the generic invalid-"
                    "credentials error.",
                    "Leaving MFA disabled until the UI lands. If an account is already enrolled, "
                    "an admin can clear it with `POST /api/users/{id}/reset-mfa`.",
                ],
                [
                    "Scanners",
                    "Only **Semgrep** and **ZAP** are launchable over HTTP. The Nuclei, "
                    "OSV-Scanner, Gitleaks, testssl, httpx and katana adapters exist in the worker "
                    "but are not exposed.",
                    "Running their `verify_*_scanner.py` scripts directly if those tools are in "
                    "scope for your release.",
                ],
                [
                    "AI agent targets",
                    "You can create one in the UI, but neither launcher accepts it — the AI/LLM "
                    "launcher lists only AI chatbot and LLM API wrapper targets.",
                    "Launching `agent_permission` over the API (§26.6).",
                ],
                [
                    "No screen yet",
                    "Log analysis, remediation generation, SARIF export/import, the triage "
                    "overview, scan-plan generation, and ROE acknowledgement history.",
                    "The API (§26).",
                ],
                [
                    "Audit log paging",
                    "The page shows the newest 100 events with no paging control.",
                    "`GET /api/audit-events?limit=500&offset=…`",
                ],
                [
                    "Health coverage",
                    "Readiness probes the database and Valkey only — the evidence store is not "
                    "checked, so an object-storage outage still shows **ok**.",
                    "Checking object storage directly when evidence misbehaves.",
                ],
                [
                    "User administration",
                    "Resetting another user's password is API-only — the UI offers it only for the "
                    "user you just created. **Deactivation cannot be undone from the product**: "
                    "there is no reactivate control and no endpoint.",
                    "`POST /api/users/{id}/reset-password`. And only ever deactivating accounts "
                    "you are willing to lose.",
                ],
                [
                    "Session invalidation",
                    "A self-service password change does not revoke your other sessions.",
                    "Using **Sign out everywhere** after changing your password.",
                ],
                [
                    "Separation of duties",
                    "An Admin can approve their own high-risk gate — there is no code check that "
                    "the approver differs from the requester.",
                    "Enforcing it as policy: Testers request, Reviewers approve.",
                ],
                [
                    "Read-only SARIF",
                    "The SARIF export route is view-gated, so Read only can pull every finding as "
                    "SARIF despite being refused report export.",
                    "Not granting Read only to anyone who should not have the findings.",
                ],
                [
                    "Evidence storage",
                    "Production evidence storage is a blocking pre-go-live gate — the development "
                    "MinIO build is not production-safe.",
                    "Confirming with your administrator which backend is in use, and that its "
                    "compliance-mode WORM has been verified.",
                ],
            ],
            (1.3, 3.0, 2.3),
        ),
    )
)

# ── 29
sec("29  Troubleshooting")
A(
    (
        "table",
        (
            ["Symptom", "Cause and fix"],
            [
                [
                    "Signed out unexpectedly mid-session",
                    "The 15-minute idle timeout, the 8-hour absolute cap, or a revocation — a role "
                    "change, a deactivation, an admin password reset, or **Sign out everywhere** "
                    "elsewhere. Sign in again.",
                ],
                [
                    "*Sign-in failed — the API is unreachable.*",
                    "Either the API really is down (check `docker compose ps` and `/health`), or "
                    "you are rate-limited after 5 failed attempts on that email. Wait out the "
                    "15-minute window.",
                ],
                [
                    "Every page bounces to `/set-password`",
                    "Your account is in forced-change state. Set a permanent password; the same "
                    "session is released immediately.",
                ],
                [
                    "`403 password_change_required` on an API call",
                    "Same cause. Only `GET /auth/me`, `POST /auth/me/password` and the logout "
                    "routes work until the password is set.",
                ],
                [
                    "`403 CSRF token missing or invalid`",
                    "You sent the session cookie without the `X-CSRF-Token` header. Take the token "
                    "from the login response and send it on every POST, PATCH and DELETE.",
                ],
                [
                    "Launch button works but the scan is refused",
                    "Read the machine reason and look it up in §17. The commonest cause by far is "
                    "an unset or expired test window.",
                ],
                [
                    "A complete engagement can never launch a scan",
                    "The test window is blank. An absent window is treated as no authorization. "
                    "Edit the engagement, set the window, re-accept the ROE.",
                ],
                [
                    "The ROE flipped back to *Acceptance required*",
                    "You edited scope, or the window, rate limit, or maximum intensity. Re-accept "
                    "it — by design.",
                ],
                [
                    "`429` with `Retry-After` when launching",
                    "The concurrency cap: 5 per engagement, 20 per organization. Wait for a running "
                    "scan to finish.",
                ],
                [
                    "`409` on an AI feature",
                    "The engagement has **Hosted LLMs allowed** unchecked and the resolved model is "
                    "hosted. Register a local model, pick it on the engagement, or tick the box if "
                    "that is authorized.",
                ],
                [
                    "`429` on an AI feature",
                    "The engagement's LLM budget ceiling is spent. That is a deliberate cost stop, "
                    "not a fault.",
                ],
                [
                    "The Framework dropdown on a finding is empty",
                    "The compliance knowledge base was never seeded. Run "
                    "`scripts/seed_compliance.py`.",
                ],
                [
                    "*The active-model status is unavailable*",
                    "The web app could not reach the API. Check the `api` container and `/health`.",
                ],
                [
                    "A URL id change shows *Not found*",
                    "Working as intended — cross-engagement access returns `404`, never data.",
                ],
                [
                    "Uploaded archive refused",
                    "Over 100 MiB → `413`. Zip-slip, zip-bomb, symlinks, or not actually an "
                    "archive → `422`. Repackage it cleanly.",
                ],
                [
                    "A target's Type dropdown is disabled",
                    "Type is immutable after creation. Delete the target and add it again.",
                ],
                [
                    "A scan is stuck",
                    "Click **Cancel**. It terminates the process group and confirms the tree is "
                    "gone. Escalate to the emergency-stop contact if the target is affected.",
                ],
            ],
            (2.4, 4.2),
        ),
    )
)

# ── 30
sec("30  Appendix A — reference tables")
A(("h2", "30.1  Engagement status"))
A(("p", "`draft` · `active` · `paused` · `closed`"))
A(("h2", "30.2  Scan intensity"))
A(
    (
        "p",
        "`passive` · `safe_active` · `authenticated_active` · `high_risk`. The launchers offer "
        "only `safe_active` and `authenticated_active`; `high_risk` requires an approval gate.",
    )
)
A(("h2", "30.3  Scope"))
A(("p", "Kind: `allow` · `deny`.  Matcher: `url` · `domain` · `ip_cidr` · `api_base` · `repo`."))
A(("h2", "30.4  Target"))
A(
    (
        "p",
        "Type: `web_app` · `rest_api` · `graphql_api` · `source_repo` · `source_archive` · "
        "`ai_chatbot` · `llm_api_wrapper` · `ai_agent`.  Environment: `dev` · `staging` · "
        "`production`.  Auth status: `none` · `configured` · `verified`.",
    )
)
A(("h2", "30.5  Scan and suite"))
A(
    (
        "p",
        "Status: `queued` · `running` · `completed` · `failed` · `cancelled`.  Suite: "
        "`prompt_injection` · `data_leakage` · `agent_permission`.  Scanner (launchable): "
        "`semgrep` · `zap`.",
    )
)
A(("h2", "30.6  Finding"))
A(
    (
        "p",
        "Severity: `critical` · `high` · `medium` · `low` · `informational`.  Provenance: "
        "`automated` · `ai_generated` · `validated` · `manually_overridden`.  Status: `open` · "
        "`in_triage` · `confirmed` · `mitigated` · `fixed` · `accepted_risk` · `false_positive` · "
        "`out_of_scope`.",
    )
)
A(("h2", "30.7  Evidence, report, approval, audit"))
A(
    (
        "p",
        "Evidence kind: `raw_scanner_output` · `http_transcript` · `llm_transcript` · "
        "`source_archive`.  Report type: `poam` · `technical` · `executive`; status: `draft` · "
        "`final`; export format: `csv` · `markdown` · `pdf` · `docx` · `json`.  Approval status: "
        "`pending` · `approved` · `denied` · `expired` · `revoked` · `consumed`.  Audit outcome: "
        "`success` · `blocked` · `failure`.  CVSS version: `v4_0` · `v3_1`.",
    )
)
A(("h2", "30.8  Defaults and limits"))
A(
    (
        "table",
        (
            ["Setting", "Default"],
            [
                ["Session idle timeout", "15 minutes, sliding"],
                ["Session absolute cap", "8 hours"],
                ["Failed sign-ins per email", "5 per 15-minute rolling window"],
                ["Failed sign-ins per IP", "30 per 15-minute rolling window"],
                ["State-changing API requests per user", "300 per minute"],
                ["Concurrent scans per engagement", "5"],
                ["Concurrent scans per organization", "20"],
                ["Engagement rate limit", "5 rps (range 1–1000)"],
                ["Minimum password length", "12 characters, and not in the breach list"],
                ["Maximum source archive", "100 MiB"],
                ["Audit events shown in the UI", "Newest 100 (API limit 500 per page)"],
                ["Recent-scans poll interval", "≈2.5 seconds while a scan is active"],
                ["Approval expiry", "Caller-set, minimum 1 hour"],
            ],
            (3.4, 3.2),
        ),
    )
)

# ── 31
sec("31  Appendix B — verification record")
A(
    (
        "p",
        "What was executed to confirm this guide describes the product as built, on "
        "**3 August 2026**, against the working tree at commit `0ded8a1` plus the in-flight AI "
        "model registry work.",
    )
)
A(("h2", "31.1  Behaviour verified against the running system"))
A(
    (
        "table",
        (
            ["Check", "Result"],
            [
                ["API unit and safety test suite (`pytest`)", "**839 passed**, 0 failed"],
                [
                    "Live readiness probe `GET /api/readyz`",
                    "`{\"status\":\"ok\",\"checks\":{\"database\":\"ok\",\"valkey\":\"ok\"}}` — "
                    "confirms the §25 gap: no evidence-store check",
                ],
                [
                    "Unauthenticated `GET /`",
                    "`307` redirect to `/login` — the auth guard described in §6 is in place",
                ],
                ["`GET /login`", "`200`, copy matches §6.1 verbatim"],
                [
                    "`GET /api/compliance/frameworks` as an admin",
                    "All **seven** frameworks with the exact versions and control counts in §21",
                ],
                [
                    "`GET /api/llm/status` and `/api/llm/models`",
                    "Registry live; a local Ollama model registered and flagged default — §15 as described",
                ],
                [
                    "`GET /api/engagements`",
                    "Returns `ai_model_id` on every engagement — the §9.2 field is real",
                ],
                [
                    "Route and capability extraction from the source",
                    "76 guarded routes mapped; §27 is generated from that map, not transcribed",
                ],
                [
                    "Enumerations, UI labels, limits and defaults",
                    "Read from the models, the label metadata and the settings object — §30 matches the code",
                ],
                [
                    "Screen-by-screen behaviour, copy and error messages",
                    "Cross-checked against the UAT execution record of 3 August 2026 "
                    "(`docs/DAS_Sentinel_UAT_Test_Script.xlsx`, evidence in `docs/uat-evidence/`), "
                    "run with Playwright against the real stack",
                ],
            ],
            (2.6, 4.0),
        ),
    )
)
A(("h2", "31.2  External standards confirmed against their publishers"))
A(
    (
        "table",
        (
            ["Standard", "Confirmed"],
            [
                [
                    "OWASP Top 10 for LLM Applications 2025",
                    "All ten codes and titles match the OWASP GenAI Security Project listing exactly, "
                    "including the 2025 changes: LLM05 *Improper Output Handling*, LLM07 *System "
                    "Prompt Leakage*, LLM08 *Vector and Embedding Weaknesses*.",
                ],
                [
                    "OWASP Top 10 for Agentic Applications 2026",
                    "All ten ASI01–ASI10 titles match the published 2026 edition exactly.",
                ],
                [
                    "OWASP Web Security Testing Guide 4.2",
                    "Still the current stable release; v5.0 remains in development. The pinned "
                    "version is correct.",
                ],
                [
                    "NIST SP 800-53 Rev 5.2.0",
                    "Released 27 August 2025 — the newest revision. Correctly pinned.",
                ],
                [
                    "NIST AI RMF 1.0 (AI 100-1) and AI 600-1 GenAI Profile",
                    "Both current; the platform references both, as it should.",
                ],
                [
                    "NIST SP 800-115 (2008)",
                    "Still the current technical testing guide; cited as-is.",
                ],
                [
                    "CVSS v4.0",
                    "Current standard, owned by FIRST.Org; five severity bands over 0.0–10.0. "
                    "v3.1 retained for historical CVEs, as the platform does.",
                ],
                ["SARIF 2.1.0", "The current OASIS standard version, which the export emits."],
            ],
            (2.4, 4.2),
        ),
    )
)
A(("h2", "31.3  Sources"))
A(
    (
        "b",
        [
            "OWASP Top 10 for LLM Applications 2025 — https://genai.owasp.org/llm-top-10/",
            "OWASP Top 10 for Agentic Applications 2026 — OWASP GenAI Security Project, announced 9 December 2025",
            "OWASP Web Security Testing Guide — https://owasp.org/www-project-web-security-testing-guide/",
            "NIST SP 800-53 Release 5.2.0 — https://csrc.nist.gov/News/2025/nist-releases-revision-to-sp-800-53-controls",
            "CVSS v4.0 specification — https://www.first.org/cvss/v4.0/specification-document",
        ],
    )
)
A(("h2", "31.4  Caveat"))
A(
    (
        "p",
        "The AI model registry (§15) and the per-engagement **AI model** field (§9.2) were "
        "uncommitted work in the tree at the time of writing. They were verified live against the "
        "running stack, but if that work changes before it is committed, §9.2, §15 and §27 are the "
        "sections to re-check.",
    )
)


# ──────────────────────────────── build ─────────────────────────────────

doc = Document()
styles(doc)
sec_ = doc.sections[0]
sec_.left_margin = sec_.right_margin = Inches(0.9)
sec_.top_margin = sec_.bottom_margin = Inches(0.85)

# Cover
t = doc.add_paragraph("DAS Sentinel", style="Title")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for text, size, bold in (
    ("User Guide", 20, True),
    ("AI security testing and automated penetration testing", 12, False),
    ("for authorized defensive assessments", 12, False),
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = ACCENT if bold else MUTED

for _ in range(6):
    doc.add_paragraph()
for text in (
    "Version 1.0  ·  3 August 2026",
    "Verified against the running application — see §31",
    "Authorized defensive assessments only",
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

doc.add_page_break()
doc.add_heading("Index", 1)
add_toc(doc)

# Footer page numbers
footer_p = doc.sections[0].footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run("DAS Sentinel User Guide  ·  page ")
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
footer_p._p.append(fld)
for r in footer_p.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED

render(doc, BLOCKS)
update_fields_on_open(doc)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("wrote", OUT)
print("blocks:", len(BLOCKS), "| routes documented:", sum(len(v) for v in grouped.values()))
