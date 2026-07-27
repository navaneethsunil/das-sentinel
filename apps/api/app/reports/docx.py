"""DOCX report exporter (M6).

Renders a report to a Word .docx from the Markdown the technical/executive
renderers already produce — sharing reports.markdown.iter_report_blocks with the
PDF exporter, so report structure has ONE source. `python-docx` is pure-Python
(MIT); unlike the PDF core fonts it handles full Unicode, so no character sanitizing
is needed. ponytail: inline formatting handles only **bold** (all our Markdown uses),
not the full inline grammar.
"""

import io

from docx import Document

from app.reports.markdown import iter_report_blocks

_HEADING_LEVEL = {"h1": 1, "h2": 2, "h3": 3}


def _add_inline(paragraph, text: str) -> None:  # noqa: ANN001 - docx Paragraph
    """Add text to a paragraph, rendering **bold** spans as bold runs."""
    for i, span in enumerate(text.split("**")):
        if span:
            paragraph.add_run(span).bold = i % 2 == 1  # odd spans are between ** pairs


def render_docx(markdown_text: str, *, title: str | None = None) -> bytes:
    """Render report Markdown to a .docx document (bytes). Pure function of the text."""
    doc = Document()
    if title:
        doc.core_properties.title = title
    for kind, _indent, text in iter_report_blocks(markdown_text):
        if kind == "blank":
            continue
        if kind in _HEADING_LEVEL:
            doc.add_heading(text, level=_HEADING_LEVEL[kind])
        elif kind == "bullet":
            _add_inline(doc.add_paragraph(style="List Bullet"), text)
        else:
            _add_inline(doc.add_paragraph(), text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
