"""M3-B5 reports — CI-safe unit tests (no DB).

Covers the pure exporters (POA&M CSV field set + formula-injection guard; Markdown
rendering) and the draft/finalized edit guard. Report assembly from live findings +
CVSS + compliance and the HTTP lifecycle are proven in scripts/verify_reports.py.
"""

import csv
import io
from datetime import UTC, datetime
from types import SimpleNamespace

import pytest

from app.models.report import ReportStatus
from app.reports.markdown import render_executive_markdown, render_markdown_report
from app.reports.poam_csv import csv_safe, render_poam_csv
from app.services.reports import ReportError, update_report

POAM_HEADERS = [
    "Weakness ID",
    "Weakness Description",
    "Affected Asset",
    "Source of Discovery",
    "Severity",
    "CVSS Score",
    "Control Mapping",
    "Recommended Remediation",
    "Responsible Owner",
    "Planned Completion Date",
    "Current Status",
    "Milestones",
    "Risk Acceptance Notes",
]


def _body() -> dict:
    return {
        "schema": "das.report/v1",
        "report_type": "poam",
        "generated_at": "2026-07-24T00:00:00+00:00",
        "engagement": {"id": "e1", "name": "Acme", "client_system_name": "Acme Portal"},
        "summary": "Overall posture is fair.",
        "findings": [
            {
                "finding_id": "f1",
                "weakness_id": "W-001",
                "title": "Prompt Injection",
                "severity": "high",
                "current_status": "open",
                "validation_status": "automated",
                "is_false_positive": False,
                "affected_asset": "chatbot (https://chat.example.com)",
                "source_of_discovery": "prompt_injection",
                "description": "The model followed injected instructions.",
                "impact": "Data exfiltration.",
                "recommended_remediation": "Add input/output guardrails.",
                "cvss": {
                    "version": "v4_0",
                    "base_score": 8.5,
                    "severity_band": "high",
                    "vector": "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:H/VA:N/SC:N/SI:N/SA:N",
                },
                "mappings": [
                    {
                        "framework_key": "owasp_llm_2025",
                        "framework_name": "OWASP LLM",
                        "code": "LLM01",
                        "title": "Prompt Injection",
                    },
                    {
                        "framework_key": "nist_800_53_r5",
                        "framework_name": "NIST 800-53",
                        "code": "SI-10",
                        "title": "Information Input Validation",
                    },
                ],
                "responsible_owner": "",
                "planned_completion_date": "",
                "milestones": "",
                "risk_acceptance_notes": "",
            }
        ],
    }


# ── CSV formula-injection guard ──────────────────────────────────────────────
@pytest.mark.parametrize("bad", ["=1+1", "+1", "-1", "@cmd", "\ttab", "\rcr"])
def test_csv_safe_neutralizes_formula_prefixes(bad: str) -> None:
    assert csv_safe(bad).startswith("'")


def test_csv_safe_leaves_safe_values() -> None:
    assert csv_safe("high") == "high"
    assert csv_safe(None) == ""
    assert csv_safe(8.5) == "8.5"


# ── POA&M CSV ────────────────────────────────────────────────────────────────
def test_poam_csv_headers_and_row() -> None:
    rows = list(csv.reader(io.StringIO(render_poam_csv(_body()))))
    assert rows[0] == POAM_HEADERS
    row = dict(zip(POAM_HEADERS, rows[1], strict=True))
    assert row["Weakness ID"] == "W-001"
    assert "Prompt Injection" in row["Weakness Description"]
    assert row["Affected Asset"] == "chatbot (https://chat.example.com)"
    assert row["Source of Discovery"] == "prompt_injection"
    assert row["Severity"] == "high"
    assert row["CVSS Score"] == "8.5 (CVSS v4.0)"
    assert row["Control Mapping"] == "LLM01 (OWASP LLM); SI-10 (NIST 800-53)"
    assert row["Recommended Remediation"] == "Add input/output guardrails."
    assert row["Current Status"] == "open"


def test_poam_csv_neutralizes_malicious_title() -> None:
    body = _body()
    body["findings"][0]["title"] = "=HYPERLINK(evil)"
    body["findings"][0]["description"] = None
    rows = list(csv.reader(io.StringIO(render_poam_csv(body))))
    row = dict(zip(POAM_HEADERS, rows[1], strict=True))
    assert row["Weakness Description"].startswith("'=HYPERLINK")


def test_poam_csv_unscored_and_unmapped() -> None:
    body = _body()
    body["findings"][0]["cvss"] = None
    body["findings"][0]["mappings"] = []
    rows = list(csv.reader(io.StringIO(render_poam_csv(body))))
    row = dict(zip(POAM_HEADERS, rows[1], strict=True))
    assert row["CVSS Score"] == ""
    assert row["Control Mapping"] == ""


# Columns the brief §15 POA&M mandates — pinned independently of the exporter's own
# _COLUMNS so a silent drop/rename of a required column fails the build (M3-T3).
_REQUIRED_POAM_COLUMNS = frozenset(
    {
        "Weakness ID",
        "Weakness Description",
        "Affected Asset",
        "Source of Discovery",
        "Severity",
        "CVSS Score",
        "Control Mapping",
        "Recommended Remediation",
        "Responsible Owner",
        "Planned Completion Date",
        "Current Status",
        "Milestones",
        "Risk Acceptance Notes",
    }
)


def test_poam_csv_has_all_required_columns() -> None:
    header = next(csv.reader(io.StringIO(render_poam_csv(_body()))))
    assert _REQUIRED_POAM_COLUMNS.issubset(header)  # no mandated column dropped
    assert len(header) == len(_REQUIRED_POAM_COLUMNS)  # no stray/duplicate columns


# ── Markdown ─────────────────────────────────────────────────────────────────
def test_markdown_contains_expected_sections() -> None:
    md = render_markdown_report(_body())
    assert "**Engagement:** Acme — Acme Portal" in md
    assert "## Summary" in md
    assert "Overall posture is fair." in md
    assert "### W-001 — Prompt Injection" in md
    assert "- **Severity:** high" in md
    assert "8.5 (CVSS v4.0)" in md
    assert "**OWASP mapping:** LLM01 Prompt Injection" in md
    assert "**NIST mapping:** SI-10 Information Input Validation" in md
    assert "The model followed injected instructions." in md


def test_markdown_poam_fields_only_when_filled() -> None:
    body = _body()
    assert "POA&M tracking" not in render_markdown_report(body)
    body["findings"][0]["responsible_owner"] = "Team Blue"
    md = render_markdown_report(body)
    assert "POA&M tracking" in md
    assert "Team Blue" in md


def test_markdown_unscored_shows_not_scored() -> None:
    body = _body()
    body["findings"][0]["cvss"] = None
    assert "Not scored" in render_markdown_report(body)


# ── edit guard ───────────────────────────────────────────────────────────────
class _FakeSession:
    def __init__(self) -> None:
        self.flushed = False

    async def flush(self) -> None:
        self.flushed = True


@pytest.mark.asyncio
async def test_update_report_draft_applies_changes() -> None:
    report = SimpleNamespace(status=ReportStatus.DRAFT, title="t", body={"a": 1}, updated_at=None)
    session = _FakeSession()
    await update_report(session, report, title="new", body={"b": 2}, now=datetime.now(UTC))  # type: ignore[arg-type]
    assert report.title == "new"
    assert report.body == {"b": 2}
    assert session.flushed is True


@pytest.mark.asyncio
async def test_update_report_finalized_is_immutable() -> None:
    report = SimpleNamespace(status=ReportStatus.FINAL, title="t", body={"a": 1}, updated_at=None)
    with pytest.raises(ReportError):
        await update_report(
            _FakeSession(),  # type: ignore[arg-type]
            report,
            title="x",
            body=None,
            now=datetime.now(UTC),
        )
    assert report.title == "t"  # unchanged


# ── Executive report (M6) ────────────────────────────────────────────────────
def test_executive_report_risk_posture_top_risks_and_coverage() -> None:
    md = render_executive_markdown(_body())
    assert "# " in md and "Executive summary" in md
    assert "Overall posture is fair." in md  # editable summary carried through
    assert "**Total findings:** 1" in md
    assert "**High:** 1" in md  # severity breakdown
    assert "Prompt Injection" in md  # top risk listed
    # compliance coverage lists both frameworks + their codes
    assert "OWASP LLM:** 1 control(s) — LLM01" in md
    assert "NIST 800-53:** 1 control(s) — SI-10" in md
    # executive view omits the per-finding deep-detail narrative
    assert "Recommended remediation" not in md


def test_executive_top_risks_ranked_by_severity_then_cvss() -> None:
    body = _body()
    findings = body["findings"]
    low = {**findings[0], "finding_id": "f2", "title": "Low Thing", "severity": "low", "cvss": None}
    crit = {
        **findings[0],
        "finding_id": "f3",
        "title": "Critical Thing",
        "severity": "critical",
        "cvss": {**findings[0]["cvss"], "base_score": 9.8, "severity_band": "critical"},
    }
    body["findings"] = [low, findings[0], crit]  # deliberately unsorted
    md = render_executive_markdown(body)
    # worst-first: critical before high before low in the top-risks section
    assert md.index("Critical Thing") < md.index("Prompt Injection") < md.index("Low Thing")
    assert "**Total findings:** 3" in md
    assert "**Critical:** 1" in md and "**High:** 1" in md and "**Low:** 1" in md


def test_executive_report_no_findings() -> None:
    body = _body()
    body["findings"] = []
    md = render_executive_markdown(body)
    assert "**Total findings:** 0" in md
    assert "_No findings._" in md
    assert "_No compliance mappings._" in md


# ── PDF (M6) ─────────────────────────────────────────────────────────────────
def test_pdf_from_markdown_is_valid_pdf() -> None:
    from app.reports.markdown import render_markdown_report
    from app.reports.pdf import render_pdf

    pdf = render_pdf(render_markdown_report(_body()), title="Acme")
    assert isinstance(pdf, bytes)
    assert pdf.startswith(b"%PDF-")  # valid PDF header
    assert pdf.rstrip().endswith(b"%%EOF")
    assert len(pdf) > 500  # non-trivial document


def test_pdf_survives_non_latin_text() -> None:
    from app.reports.pdf import render_pdf

    # core fonts are Latin-1; non-Latin must degrade, never crash the export
    pdf = render_pdf("# Rapport\n\n- **Finding:** 日本語 مرحبا café", title="x")
    assert pdf.startswith(b"%PDF-")


# ── DOCX (M6) ────────────────────────────────────────────────────────────────
def test_docx_is_valid_and_carries_content() -> None:
    import io as _io

    from docx import Document

    from app.reports.docx import render_docx
    from app.reports.markdown import render_markdown_report

    blob = render_docx(render_markdown_report(_body()), title="Acme")
    assert isinstance(blob, bytes)
    assert blob.startswith(b"PK\x03\x04")  # docx is a zip archive
    doc = Document(_io.BytesIO(blob))  # re-opens = well-formed
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Prompt Injection" in text
    assert "Severity:" in text  # **bold** run flattened into paragraph text


def test_docx_preserves_unicode() -> None:
    import io as _io

    from docx import Document

    from app.reports.docx import render_docx

    blob = render_docx("# R\n\n- **F:** 日本語 café", title="x")
    doc = Document(_io.BytesIO(blob))
    assert any("日本語" in p.text for p in doc.paragraphs)  # full unicode, no sanitizing
